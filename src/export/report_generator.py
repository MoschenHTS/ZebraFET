import io
import os
import logging
import tempfile
from collections import Counter
from datetime import datetime
from typing import Dict, Any, List, Tuple, Union, Optional

import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.table import Table
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from PIL import Image, ImageDraw, ImageFont, ImageOps

from src.core.constants import APP_VERSION, HATCHED_STATUSES, LIVE_STATUSES, DEFAULT_SPECIES
from src.core.utils import resource_path
from src.core.biostatistics import (wilson_ci, odds_ratio_ci, significance_marker,
                                    format_with_unit, format_ci_range,
                                    evaluable_n, select_control_rows,
                                    mortality_bounding_concentrations,
                                    CONTROL_MODE_POOLED, CONTROL_MODE_NEGATIVE,
                                    CONTROL_MODE_SOLVENT)

log = logging.getLogger(__name__)

# Style Constants 
FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = 12
FONT_SIZE_TABLE = 10
FONT_SIZE_TITLE = 14
MARGIN_CM = 2.54

class ReportGenerator:
    """
    Generates a publication-ready Word document report for a FET test,
    closely following the OECD TG 236 guidelines.
    """

    #: Sections the operator may omit, in the order they appear. Test validity and
    #: the lethal endpoints are not listed: a FET report without them is not a
    #: report, so they are always emitted. References are likewise always emitted —
    #: the methods text cites its sources inline, and dropping the list would leave
    #: those citations unresolved and remove ZebraFET's own citation with them.
    OPTIONAL_SECTIONS = (
        ("methods", "Materials and Methods"),
        ("results_timeseries", "LC50 time-series"),
        ("results_timecourse", "Mortality time-course"),
        ("results_fate", "Fate composition"),
        ("results_sublethal", "Sublethal endpoint analysis"),
        ("results_hatching", "Hatching rate"),
        ("results_malformation", "Malformation profile"),
        ("appendix_layout", "Appendix: plate layout"),
        ("appendix_raw", "Appendix: raw observational data"),
        ("appendix_malformation", "Appendix: malformation frequency"),
        ("appendix_photos", "Appendix: photographic documentation"),
    )
    ALL_SECTIONS = frozenset(key for key, _ in OPTIONAL_SECTIONS)

    def __init__(self, snapshot: Dict[str, Any], project_dir: str,
                 analysis_results: Dict[str, Any],
                 sections: Optional[set] = None) -> None:
        self.project_data = snapshot
        self.project_dir = project_dir
        self.analysis_results = analysis_results
        #: None means every section, so existing callers are unaffected.
        self.sections = set(self.ALL_SECTIONS if sections is None else sections)
        self.document: Document = Document()
        self.figure_count: int = 0
        self.table_count: int = 0
        #: Set by generate_report; the raw-data CSV is written next to it.
        self._output_path: Optional[str] = None
        #: Keys of works actually cited, populated as the methods text is built.
        self._cited: set = set()
        #: Top-level headings are numbered as they are emitted rather than
        #: hardcoded, so omitting one does not misnumber everything after it.
        self._section_no = 0
        self._setup_document_properties()

    def _wants(self, key: str) -> bool:
        return key in self.sections

    def _begin_section(self, title: str) -> int:
        """Emit a numbered top-level heading and return its number."""
        self._section_no += 1
        self.document.add_heading(f"{self._section_no}. {title}", level=1)
        return self._section_no

    def _setup_document_properties(self) -> None:
        """Sets up the basic properties of the Word document, like margins and fonts."""
        for section in self.document.sections:
            section.top_margin = Cm(MARGIN_CM)
            section.bottom_margin = Cm(MARGIN_CM)
            section.left_margin = Cm(MARGIN_CM)
            section.right_margin = Cm(MARGIN_CM)

        style = self.document.styles['Normal']
        style.font.name = FONT_NAME
        style.font.size = Pt(FONT_SIZE_BODY)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def generate_report(self, output_path: str) -> bool:
        """
        Orchestrates the generation of the full report by calling section-specific methods.
        """
        current_section = "Initialization"
        self._output_path = output_path
        try:
            current_section = "Title Page"
            self._create_title_page()
            
            if self._wants("methods"):
                current_section = "Materials and Methods"
                self._create_materials_and_methods()

            current_section = "Results"
            self._create_results_section()
            
            # References come after Results so that every conditional methods
            # sentence has run and self._cited is complete.
            current_section = "References"
            self._create_references_section()

            current_section = "Appendices"
            self._create_appendix_section()

            self.document.save(output_path)
            log.info(f"Report successfully generated at: {output_path}")
            return True
        except Exception as e:
            log.error(f"Failed to generate report at section '{current_section}': {e}", exc_info=True)
            return False
        finally:
            plt.close('all')

    # Title Page #
    def _create_title_page(self) -> None:
        """Creates the main title page of the report."""
        substance = self.project_data.get('substance') or '[Insert Substance Name]'
        title_text = f"Fish Embryo Acute Toxicity (FET) Test (OECD TG 236) Final Report: Toxicity Assessment of {substance}"
        p_title = self.document.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run(title_text)
        run_title.font.name = FONT_NAME
        run_title.font.size = Pt(FONT_SIZE_TITLE)
        run_title.bold = True
        p_title.paragraph_format.space_after = Pt(24)
        p_title.paragraph_format.line_spacing = 1.0

        author = self.project_data.get('main_researcher') or '[Insert Researcher Name]'
        p_author = self.document.add_paragraph(f"Principal Investigator: {author}")
        p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER

        try:
            date_str = self.project_data.get("start_date")
            report_date = datetime.fromisoformat(date_str).strftime('%B %d, %Y')
        except (ValueError, TypeError):
            report_date = "[Insert Completion Date]"
        p_date = self.document.add_paragraph(f"Date of Study Completion: {report_date}")
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER

        generation_date = datetime.now().strftime('%B %d, %Y')
        p_gen = self.document.add_paragraph(f"Report Generated: {generation_date}")
        p_gen.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.document.add_page_break()

    # Materials & Methods #
    def _create_materials_and_methods(self) -> None:
        """Creates the comprehensive Materials and Methods section with clear subdivisions."""
        methods_no = self._begin_section("Materials and Methods")
        
        self.document.add_heading(f"{methods_no}.1 Test Substance", level=2)
        substance_details = self.project_data.get("substance_details", {})
        details = {
            "Common Name": self.project_data.get("substance") or "[Not specified]",
            "CAS Number": substance_details.get("cas_number") or "[Not specified]",
            "IUPAC Name": substance_details.get("iupac_name") or "[Not specified]",
            "Molecular Weight (g/mol)": substance_details.get("molecular_weight") or "[Not specified]",
            "Purity (%)": substance_details.get("purity") or "[Not specified]",
            "Water Solubility": substance_details.get("water_solubility") or "[Not specified]",
            "Supplier": substance_details.get("supplier") or "[Not specified]",
            "Physical Appearance": substance_details.get("physical_appearance") or "[Not specified]",
        }
        self._add_key_value_table("Physicochemical properties of the test substance.", details)
        
        self.document.add_heading(f"{methods_no}.2 Test Organism", level=2)
        organisms = self.project_data.get("test_organisms", {})
        org_details = {
            "Species": organisms.get("species") or DEFAULT_SPECIES,
            "Strain": organisms.get("strain") or "[Not specified]",
            "Source of Brood Stock": organisms.get("source") or "[Not specified]",
        }
        self._add_key_value_table("Details of the test organism.", org_details)
        p = self.document.add_paragraph()
        p.add_run("Egg Collection and Maintenance: ").bold = True
        p.add_run(organisms.get('collection_method') or "[Not specified]")

        self.document.add_heading(f"{methods_no}.3 Test Conditions", level=2)
        conditions = self.project_data.get("test_conditions", {})
        cond_details = {
            "Dilution Medium": conditions.get("water_type") or "[Not specified]",
            "Temperature (°C)": conditions.get("temperature") or "[Insert temperature]",
            "pH": conditions.get("ph") or "[Insert pH]",
            "Total Hardness (mg/L CaCO₃)": conditions.get("hardness") or "[Insert hardness]",
            "Conductivity (µS/cm)": conditions.get("conductivity") or "[Insert conductivity]",
            "Dissolved Oxygen": conditions.get("dissolved_oxygen") or "[Insert DO]",
            "Photoperiod (L:D)": conditions.get("photoperiod") or "[Insert photoperiod]",
            "Acceptable Control Mortality (%)": conditions.get("acceptable_mortality", "[Not specified]"),
        }
        # Omitted rather than shown as unrecorded: it is optional, and a blank row
        # would invite the reader to treat its absence as a finding.
        if conditions.get("fertilization_rate"):
            cond_details["Batch Fertilization Rate"] = conditions["fertilization_rate"]
        self._add_key_value_table("Summary of test conditions.", cond_details)
        self._add_water_quality_table()

        self.document.add_heading(f"{methods_no}.4 Experimental Design and Exposure", level=2)
        self.document.add_heading("Test Concentrations and Controls", level=3)
        concentrations = self.project_data.get("concentration_settings", {}).get("concentrations", [])
        substrate_concs = [str(c['value']) for c in concentrations if c['type'] == 'Substrate']
        conc_text = ", ".join(substrate_concs) if substrate_concs else "[Not specified]"
        conc_unit = self.project_data.get("concentration_unit", "unit")
        self.document.add_paragraph(f"{len(substrate_concs)} nominal concentrations of the test substance were used: {conc_text} {conc_unit}.")
        
        solvent = self.project_data.get("substance_details", {}).get("solvent_used")
        if solvent:
             self.document.add_paragraph(f"Both a negative control (dilution water only) and a solvent control ({solvent or '[Not specified]'}) were run in parallel.")
        else:
             self.document.add_paragraph(f"A negative control (dilution water only) was run in parallel.")

        positive_controls = [c for c in concentrations if c['type'] == 'Positive Control']
        if positive_controls:
            pc_substance = self.project_data.get("substance_details", {}).get("positive_control_substance", "[Not specified]")
            pc_conc = positive_controls[0]['value']
            self.document.add_paragraph(
                f"A positive control using {pc_substance} at a concentration of {pc_conc} {conc_unit} was included to validate test sensitivity."
            )

        self.document.add_heading("Test Procedure and Exposure System", level=3)
        if concentrations:
            num_replicates = concentrations[0].get('replicates', '[N/A]')
            embryos_per_well = concentrations[0].get('wells', 1)
            rep_word = "replicate" if num_replicates == 1 else "replicates"
            emb_word = "embryo" if embryos_per_well == 1 else "embryos"
            design_text = (f"The experimental design consisted of {num_replicates} {rep_word} for each treatment group. "
                           f"Each replicate comprised {embryos_per_well} {emb_word} per well.")
            self.document.add_paragraph(design_text)

        plate_format = self.project_data.get("plate_format", "96-well")
        from src.core.constants import PLATE_FORMATS
        p_rows, p_cols = PLATE_FORMATS.get(plate_format, (8, 12))
        p = self.document.add_paragraph()
        p.add_run("Test Container: ").bold = True
        p.add_run(
            f"{plate_format} multi-well plates "
            f"({p_rows}\u00d7{p_cols} grid, {p_rows * p_cols} wells per plate) were used."
        )

        methodology = self.project_data.get("methodology", {})
        test_procedure = methodology.get('test_procedure', '[Not specified]')
        renewal_text = "Solutions were renewed every 24 hours." if "renewal" in test_procedure else "No solution renewal was performed."
        p = self.document.add_paragraph()
        p.add_run("Exposure Type: ").bold = True
        p.add_run(f"The test was conducted under {test_procedure} conditions. {renewal_text}")

        p = self.document.add_paragraph()
        p.add_run("Solution Preparation: ").bold = True
        p.add_run(methodology.get('solution_preparation') or "[Not specified]")
             
        self.document.add_heading(f"{methods_no}.5 Endpoints Assessed", level=2)
        self.document.add_paragraph(
            f"Endpoints were assessed daily up to {(self.project_data.get('num_days') or 4) * 24} hours "
            "post-fertilization (hpf) according to OECD TG 236 criteria. "
            "Lethal endpoints included coagulation of the embryo, lack of somite formation, non-detachment of the tail, and absence of heartbeat. "
            "Sub-lethal endpoints observed included hatching rate and the incidence of morphological abnormalities (e.g., pericardial edema, notochord deformities, body curvature). "
            "Morphological abnormalities are expressed as a percentage of surviving embryos, as coagulated or dead embryos cannot be scored for malformation."
        )

        self.document.add_heading(f"{methods_no}.6 Statistical Analysis and Validity Criteria", level=2)
        lc50_results = self.analysis_results.get("lc50_results", {})
        model_info = lc50_results.get("model_info", {})
        lc50_model_sentence = self._build_lc50_methods_sentence(
            model_info, lc50_results.get("bootstrap_method")
        )
        self._cite("oecd236", "fisher")
        abbott_sentence = ""
        if self.analysis_results.get("abbott_applied"):
            self._cite("abbott")
            abbott_sentence = (
                "Prior to curve fitting, mortality data were adjusted for control (background) mortality "
                "using Abbott's formula (Abbott, 1925). "
            )
        # 'available' is true whenever any sublethal output exists, including the
        # pooled ">= 1 abnormality" endpoint on its own. The per-endpoint battery
        # only runs when at least one morphological endpoint was scored, so the
        # methods text must be gated on the tests themselves rather than the flag.
        sublethal_stats = self.analysis_results.get("sublethal_stats", {}) or {}
        sublethal_sentence = ""
        if sublethal_stats.get("tests"):
            self._cite("benjamini_hochberg", "armitage", "fisher")
            sublethal_sentence = (
                "Each sublethal (teratogenic) endpoint was compared with the control by Fisher's Exact Test "
                "with Benjamini-Hochberg control of the false discovery rate (Benjamini & Hochberg, 1995), and "
                "per-endpoint dose-related trends were assessed by the Cochran-Armitage test. "
            )
        ti_sentence = (
            "A teratogenic index (TI = LC50 / EC50) was derived from the malformation dose-response to "
            "distinguish selective developmental toxicity from general lethality. "
            if self.analysis_results.get("teratogenic_index", {}).get("ti_numeric") is not None else ""
        )
        # Only claim the trend test / effect sizes when they were actually computed
        # (both require a control group; the trend test also needs >= 2 dose groups).
        trend_computed = self.analysis_results.get("trend_results", {}).get("p_value") not in (None, "Not Calculated")
        summary_df = self.analysis_results.get("summary_df", pd.DataFrame())
        has_control = (
            summary_df is not None and not summary_df.empty and "conc_type" in summary_df.columns
            and not select_control_rows(summary_df, self._control_mode()).empty
        )
        trend_sentence = ""
        if trend_computed:
            self._cite("armitage")
            trend_sentence = (
                "A Cochran-Armitage test for trend was additionally applied to assess a monotonic dose-related "
                "increase in mortality (Armitage, 1955). "
            )
        effect_sentence = ""
        if has_control:
            self._cite("wilson", "haldane", "anscombe")
            effect_sentence = (
                "Effect sizes are reported as odds ratios with 95% confidence intervals (Haldane-Anscombe "
                "corrected for zero cells; Haldane, 1956; Anscombe, 1956), and group proportions are accompanied "
                "by Wilson score 95% confidence intervals (Wilson, 1927). "
            )
        self.document.add_paragraph(
            "The test was conducted following the OECD Guideline for the Testing of Chemicals, No. 236: Fish Embryo Acute Toxicity (FET) Test (OECD, 2025). "
            "Data acquisition and analysis were performed using the ZebraFET — Standardized Zebrafish Embryo Toxicity Test Assistant software. "
            f"{self._validity_criteria_sentence()}"
            f"Statistical analyses were performed using Fisher's Exact Test (Fisher, 1922) with {self._noec_correction_label()} correction for multiple comparisons to determine the No Observed Effect Concentration (NOEC) and Lowest Observed Effect Concentration (LOEC). "
            f"{self._control_reference_sentence()}"
            f"{trend_sentence}"
            f"{sublethal_sentence}"
            f"{effect_sentence}"
            f"{abbott_sentence}"
            f"{ti_sentence}"
            f"{lc50_model_sentence}"
            f"{self._weighting_sentence()}"
            f"{self._tsk_methods_sentence()}"
        )

    # References #
    #: Works the report may cite, keyed by the tag used at each citation site.
    #: Only those recorded in self._cited are listed, so the section can never
    #: credit a method the analysis did not actually apply.
    _REFERENCES = {
        "abbott": (
            "Abbott, W. S. (1925). A method of computing the effectiveness of an "
            "insecticide. Journal of Economic Entomology, 18(2), 265–267."
        ),
        "armitage": (
            "Armitage, P. (1955). Tests for linear trends in proportions and "
            "frequencies. Biometrics, 11(3), 375–386."
        ),
        "benjamini_hochberg": (
            "Benjamini, Y., & Hochberg, Y. (1995). Controlling the false discovery rate: "
            "a practical and powerful approach to multiple testing. Journal of the Royal "
            "Statistical Society: Series B, 57(1), 289–300."
        ),
        "hamilton": (
            "Hamilton, M. A., Russo, R. C., & Thurston, R. V. (1977). Trimmed "
            "Spearman-Karber method for estimating median lethal concentrations in "
            "toxicity bioassays. Environmental Science & Technology, 11(7), 714–719."
        ),
        "holm": (
            "Holm, S. (1979). A simple sequentially rejective multiple test procedure. "
            "Scandinavian Journal of Statistics, 6(2), 65–70."
        ),
        "burnham_anderson": (
            "Burnham, K. P., & Anderson, D. R. (2002). Model Selection and Multimodel "
            "Inference: A Practical Information-Theoretic Approach (2nd ed.). Springer."
        ),
        "oecd236": (
            "OECD (2025). Test No. 236: Fish Embryo Acute Toxicity (FET) Test. OECD "
            "Guidelines for the Testing of Chemicals, Section 2. OECD Publishing, Paris."
        ),
        "anscombe": (
            "Anscombe, F. J. (1956). On estimating binomial response relations. "
            "Biometrika, 43(3–4), 461–464."
        ),
        "dunn": (
            "Dunn, O. J. (1961). Multiple comparisons among means. Journal of the "
            "American Statistical Association, 56(293), 52–64."
        ),
        "efron": (
            "Efron, B. (1979). Bootstrap methods: another look at the jackknife. The "
            "Annals of Statistics, 7(1), 1–26."
        ),
        "fisher": (
            "Fisher, R. A. (1922). On the interpretation of χ² from contingency tables, "
            "and the calculation of P. Journal of the Royal Statistical Society, 85(1), "
            "87–94."
        ),
        "haldane": (
            "Haldane, J. B. S. (1956). The estimation and significance of the logarithm "
            "of a ratio of frequencies. Annals of Human Genetics, 20(4), 309–311."
        ),
        "ritz": (
            "Ritz, C. (2010). Toward a unified approach to dose–response modeling in "
            "ecotoxicology. Environmental Toxicology and Chemistry, 29(1), 220–229."
        ),
        "wilson": (
            "Wilson, E. B. (1927). Probable inference, the law of succession, and "
            "statistical inference. Journal of the American Statistical Association, "
            "22(158), 209–212."
        ),
    }

    def _cite(self, *keys: str) -> None:
        """Record that the report relies on these works."""
        self._cited.update(k for k in keys if k in self._REFERENCES)

    def _zebrafet_citation(self) -> str:
        """Self-citation, versioned so it cannot drift from the running build."""
        return (
            "Moschen, H. T. S., Zimerman, J., Matos, V. F., Liberal, C. H. C., "
            "Bellozi, P. M. Q., de Bem, A. F., & Goulart, J. T. (2026). ZebraFET — "
            f"Standardized Zebrafish Embryo Toxicity Test Assistant (OECD TG 236) "
            f"(v{APP_VERSION}). Zenodo. https://doi.org/10.5281/zenodo.20183712"
        )

    def _create_references_section(self) -> None:
        """List the works cited in the methods, alphabetically."""
        entries = sorted(
            [self._REFERENCES[key] for key in self._cited] + [self._zebrafet_citation()]
        )
        self._begin_section("References")
        for entry in entries:
            p = self.document.add_paragraph(entry, style='Normal')
            p.paragraph_format.space_after = Pt(6)

    #: Fixed OECD TG 236 §9 thresholds. Hard-coded like the water-quality ranges
    #: above: they are properties of the guideline, whereas acceptable_mortality is
    #: a threshold laboratories legitimately tighten and so stays configurable.
    CONTROL_HATCHING_MINIMUM_PCT = 80.0            # §9e
    BATCH_FERTILISATION_MINIMUM_PCT = 70.0         # §9a
    POSITIVE_CONTROL_MINIMUM_MORTALITY_PCT = 30.0  # §9d

    # Results #
    @staticmethod
    def get_test_validity_message(
        summary_df: pd.DataFrame,
        project_data: Dict,
        control_mode: str = CONTROL_MODE_POOLED,
        end_of_test: bool = False,
        timepoint: str = "the end of the test",
        water_quality_findings: Optional[List[Dict[str, Any]]] = None,
    ) -> Optional[str]:
        """State which OECD TG 236 §9 validity criteria the test met, or None.

        Each criterion is reported only when its inputs exist. A criterion the
        software was never given data for is not a finding about the study, so it
        is omitted rather than announced as unevaluated; when nothing at all can be
        assessed this returns None and the caller drops the section.

        The verdict is confined to the criteria actually assessed. Drawn from
        control mortality alone it claimed more than had been checked, passing a
        test whose dissolved oxygen never reached the guideline minimum.

        Hatching (§9e) and positive-control sensitivity (§9d) are stated by the
        guideline at the end of the 96 hrs exposure, so both are assessed only when
        *end_of_test* is set; embryos hatch from roughly 48-72 hpf and an interim
        timepoint would fail the hatching criterion in every study.
        """
        conditions = project_data.get("test_conditions", {}) or {}
        controls = (
            select_control_rows(summary_df, control_mode)
            if not summary_df.empty else summary_df
        )
        n_control = int(evaluable_n(controls).sum()) if not controls.empty else 0

        criteria = [
            ReportGenerator._criterion_fertilization(conditions),
            ReportGenerator._criterion_control_mortality(
                controls, n_control, conditions.get("acceptable_mortality", 10.0)
            ),
            ReportGenerator._criterion_control_hatching(
                controls, n_control, end_of_test, timepoint
            ),
            ReportGenerator._criterion_positive_control(
                summary_df, end_of_test, timepoint
            ),
        ]
        criteria += [
            ReportGenerator._criterion_water_quality(f)
            for f in water_quality_findings or []
        ]

        assessed = [text for ok, text in filter(None, criteria) if ok]
        failures = [text for ok, text in filter(None, criteria) if not ok]
        total = len(assessed) + len(failures)
        if not total:
            return None

        if failures:
            # Only the positive claim needed softening; a failure must be
            # impossible to miss, so it stays blunt.
            verdict = (
                f"The test is INVALID against the OECD TG 236 validity criteria assessed "
                f"here, meeting {len(assessed)} of {total}. It does not meet the following: "
                + "; ".join(failures) + "."
            )
            if assessed:
                verdict += " The criteria it does meet are: " + "; ".join(assessed) + "."
            return verdict

        scope = ("the one OECD TG 236 validity criterion" if total == 1
                 else f"all {total} OECD TG 236 validity criteria")
        return f"The test meets {scope} assessed here: " + "; ".join(assessed) + "."

    @staticmethod
    def _criterion_fertilization(conditions: Dict) -> Optional[Tuple[bool, str]]:
        """TG 236 §9a — overall fertilization rate of the batch, >= 70%."""
        rate = ReportGenerator._leading_float(conditions.get("fertilization_rate"))
        if rate is None:
            return None
        minimum = ReportGenerator.BATCH_FERTILISATION_MINIMUM_PCT
        return (
            rate >= minimum,
            f"the batch fertilization rate was {rate:g}%, "
            + ("at or above" if rate >= minimum else "below")
            + f" the minimum of {minimum:.0f}%",
        )

    @staticmethod
    def _criterion_control_mortality(
        controls: pd.DataFrame, n_control: int, threshold: float,
    ) -> Optional[Tuple[bool, str]]:
        """TG 236 §9c — control survival >= 90%, expressed here as mortality.

        Judged against the same reference the dose comparisons use, so the verdict
        cannot disagree with the NOEC about the baseline.
        """
        if n_control <= 0:
            return None
        mortality = controls["dead"].sum() / n_control * 100
        return (
            mortality <= threshold,
            f"mortality in the control group(s) was {mortality:.2f}%, "
            + ("within" if mortality <= threshold else "exceeding")
            + f" the acceptable limit of {threshold:.0f}%",
        )

    @staticmethod
    def _criterion_control_hatching(
        controls: pd.DataFrame, n_control: int, end_of_test: bool, timepoint: str,
    ) -> Optional[Tuple[bool, str]]:
        """TG 236 §9e — control hatching rate >= 80% at the end of the exposure.

        'hatched' counts live and dead hatched larvae alike over the embryos
        actually scored, the same quantity the summary table reports as
        "Hatched (%)", so the verdict and the table cannot quote different rates.
        """
        if n_control <= 0 or not end_of_test or "hatched" not in controls.columns:
            return None
        rate = controls["hatched"].sum() / n_control * 100
        minimum = ReportGenerator.CONTROL_HATCHING_MINIMUM_PCT
        return (
            rate >= minimum,
            f"the control hatching rate at {timepoint} was {rate:.2f}%, "
            + ("at or above" if rate >= minimum else "below")
            + f" the minimum of {minimum:.0f}%",
        )

    @staticmethod
    def _criterion_positive_control(
        summary_df: pd.DataFrame, end_of_test: bool, timepoint: str,
    ) -> Optional[Tuple[bool, str]]:
        """TG 236 §9d — positive control mortality >= 30% at the end of the exposure.

        Many FET studies run without a positive control; such a design yields no
        criterion at all rather than an unmet one. Raw counts are used, as the
        guideline specifies: this group is never background-corrected.
        """
        if summary_df.empty or not end_of_test or "conc_type" not in summary_df.columns:
            return None
        group = summary_df[summary_df["conc_type"] == "Positive Control"]
        if group.empty:
            return None
        n = int(evaluable_n(group).sum())
        if n <= 0:
            return None
        mortality = group["dead"].sum() / n * 100
        minimum = ReportGenerator.POSITIVE_CONTROL_MINIMUM_MORTALITY_PCT
        return (
            mortality >= minimum,
            f"mortality in the positive control at {timepoint} was {mortality:.2f}%, "
            + ("at or above" if mortality >= minimum else "below")
            + f" the minimum of {minimum:.0f}%",
        )

    @staticmethod
    def _criterion_water_quality(finding: Dict[str, Any]) -> Tuple[bool, str]:
        """TG 236 §9b and §9f, from the per-day monitoring log."""
        if finding["passed"]:
            return True, f"{finding['subject'].lower()} remained within {finding['criterion']}"
        return False, (
            f"{finding['subject'].lower()} fell outside {finding['criterion']} on "
            + ", ".join(finding["out_of_range"])
        )

    #: Human-readable description of each reference-control selection, used in the
    #: methods text so a reader can tell which baseline the endpoints refer to.
    _CONTROL_MODE_DESCRIPTIONS = {
        CONTROL_MODE_POOLED: "the pooled control (negative and solvent controls combined)",
        CONTROL_MODE_NEGATIVE: "the negative (dilution water) control",
        CONTROL_MODE_SOLVENT: "the solvent control",
    }

    def _control_mode(self) -> str:
        """The reference control the analysis was run against."""
        return self.analysis_results.get("control_mode") or CONTROL_MODE_POOLED

    def _control_reference_sentence(self) -> str:
        """State the reference control, and the evidence behind choosing it.

        The reference determines the NOEC/LOEC, Abbott's correction, the odds
        ratios and the validity verdict alike, so a report that omits it is not
        reproducible. Where both a negative and a solvent control exist, the
        comparison between them is reported too, since that is what justifies
        pooling them or keeping them apart.
        """
        description = self._CONTROL_MODE_DESCRIPTIONS.get(
            self._control_mode(), self._CONTROL_MODE_DESCRIPTIONS[CONTROL_MODE_POOLED]
        )
        sentence = f"All comparisons were made against {description}. "
        comparison = self.analysis_results.get("control_comparison") or {}
        if comparison.get("applicable") and comparison.get("p_value") is not None:
            negative, solvent = comparison["negative"], comparison["solvent"]
            verdict = ("differed significantly" if comparison.get("differ")
                       else "did not differ significantly")
            sentence += (
                f"Mortality in the negative control ({negative['dead']}/{negative['n']}; "
                f"{negative['pct']:.1f}%) and the solvent control ({solvent['dead']}/"
                f"{solvent['n']}; {solvent['pct']:.1f}%) {verdict} "
                f"(Fisher's Exact Test, two-sided, p = {comparison['p_value']:.3g}). "
            )
        return sentence

    def _add_unevaluable_groups_note(self) -> None:
        """Record any group the statistics excluded for want of a scored embryo.

        Such a group still appears in the summary table with its assigned wells,
        so without this the reader has no way to tell that the endpoints rest on
        a smaller design than the one described in the methods.
        """
        groups = self.analysis_results.get("unevaluable_groups") or []
        if not groups:
            return
        names = ", ".join(str(g) for g in groups)
        plural = "group" if len(groups) == 1 else "groups"
        was = "was" if len(groups) == 1 else "were"
        self.document.add_paragraph(
            f"No embryos were scored in {plural} {names}, which {was} therefore excluded "
            "from the statistical analysis. The endpoints reported below are based on the "
            "remaining concentration groups.",
            style='Normal',
        )

    def _analysis_day(self) -> int:
        """The day the reported endpoints were computed for.

        The analysis worker publishes it; fall back to the project's configured
        duration and finally to the OECD default of four days, so a results dict
        produced before this field existed cannot break report generation.
        """
        day = self.analysis_results.get("analysis_day")
        if not day:
            day = self.project_data.get("num_days") or 4
        return int(day)

    def _analysis_hpf(self) -> int:
        """Hours post-fertilization represented by the analysed day."""
        return self._analysis_day() * 24

    def _is_end_of_test(self) -> bool:
        """True when the analysed day is the last day of the test.

        TG 236 states the hatching and positive-control criteria at the end of the
        96 hrs exposure, so both are judged only once the analysed day has reached
        the test's final day.
        """
        num_days = self.project_data.get("num_days")
        try:
            return num_days is not None and self._analysis_day() >= int(num_days)
        except (TypeError, ValueError):
            return False

    def _timepoint_list(self) -> str:
        """Observation timepoints up to the analysed day, as prose."""
        last = self._analysis_hpf() // 24
        points = [f"{d * 24}" for d in range(1, last + 1)]
        if len(points) == 1:
            return f"{points[0]} hpf"
        return ", ".join(points[:-1]) + f" and {points[-1]} hpf"

    def _create_results_section(self) -> None:
        """Creates the Results section, including tables and plots."""
        self._results_no = self._begin_section("Results")
        
        summary_df = self.analysis_results.get("summary_df", pd.DataFrame())
        validity_text = self.get_test_validity_message(
            summary_df, self.project_data, self._control_mode(),
            end_of_test=self._is_end_of_test(),
            timepoint=f"{self._analysis_hpf()} hpf",
            water_quality_findings=self._water_quality_findings(),
        )
        water_notes = self._water_quality_validity_notes()
        plate_notes = self._internal_plate_control_notes()

        # Subsections are numbered from a running counter: validity is emitted only
        # when there is something to report, and the headings must stay contiguous.
        self._results_sub = 0
        if (validity_text or water_notes or plate_notes
                or self.analysis_results.get("unevaluable_groups")):
            self._results_sub += 1
            self.document.add_heading(
                f"{self._results_no}.{self._results_sub} Test Validity Criteria", level=2
            )
            if validity_text:
                self.document.add_paragraph(validity_text, style='Normal')
            self._add_unevaluable_groups_note()
            for note in water_notes + plate_notes:
                self.document.add_paragraph(note, style='Normal')

        self._results_sub += 1
        self.document.add_heading(
            f"{self._results_no}.{self._results_sub} Lethal Effects and Calculated Endpoints",
            level=2,
        )
        self.document.add_paragraph(
            f"Cumulative mortality and other sublethal effects were recorded at {self._analysis_hpf()} hours post-fertilization. "
            "The data are summarized in the table below."
        )
        try:
            self._add_results_summary_table(summary_df)
        except Exception as e:
            log.warning(f"Could not create results summary table: {e}")
            self.document.add_paragraph(
                "The results summary table could not be generated due to missing or corrupt data.", style='Normal'
            )

        try:
            self._add_mortality_effects_table(summary_df)
        except Exception as e:
            log.warning(f"Could not create mortality effects table: {e}")

        conc_unit = self.project_data.get("concentration_unit", "unit")
        lc50_results = self.analysis_results.get("lc50_results", {})
        lc50 = lc50_results.get("lc50", "not calculated")
        slope = lc50_results.get("slope", "not calculated")
        r_squared = lc50_results.get("r_squared", "N/A")
        noec_loec = self.analysis_results.get("noec_loec_results", {})
        noec = noec_loec.get("noec", "not calculated")
        loec = noec_loec.get("loec", "not calculated")
        # Whether the LC50 is a number is decided by the fit, never by the leading
        # character of the display string: the failure message "100% mortality at
        # all concentrations..." begins with a digit and would otherwise be
        # rendered as though it were a concentration.
        lc50_is_numeric = bool(lc50_results.get("_fitted_params"))
        noec_text = format_with_unit(noec_loec.get("noec_numeric"), noec, conc_unit)
        loec_text = format_with_unit(noec_loec.get("loec_numeric"), loec, conc_unit)
        abbott_note = (
            " The LC50 was derived from mortality adjusted for control mortality using Abbott's "
            "correction; both observed and Abbott-corrected mortality are given in the summary table."
            if self.analysis_results.get("abbott_applied") else ""
        )
        hpf = self._analysis_hpf()
        if lc50_is_numeric:
            self.document.add_paragraph(
                f"Based on the {hpf}-hour mortality data, the LC50 was determined to be {lc50} {conc_unit} "
                f"(slope: {slope}; R\u00b2 = {r_squared}). "
                f"The No Observed Effect Concentration (NOEC) was {noec_text}, "
                f"and the Lowest Observed Effect Concentration (LOEC) was {loec_text}.{abbott_note}"
            )
        else:
            # Several LC50 failure messages already end in a full stop; adding
            # another would render as "...concentration tested.. The NOEC was".
            lc50_clause = lc50 if str(lc50).endswith(".") else f"{lc50}."
            self.document.add_paragraph(
                f"LC50 calculation: {lc50_clause} "
                f"The No Observed Effect Concentration (NOEC) was {noec_text}, "
                f"and the Lowest Observed Effect Concentration (LOEC) was {loec_text}.{abbott_note}"
            )

        self._add_mortality_bounds_sentence()
        self._add_tsk_sentence()
        self._add_trend_test_sentence()
        self._add_teratogenic_index_sentence()
        self._add_curve_fitting_section()

        mortality_fig = self.analysis_results.get("mortality_plot_figure")
        if mortality_fig:
            self._add_plot_to_document(
                mortality_fig,
                caption=f"Concentration-response curve for mortality of {self._species()} embryos after {self._analysis_hpf()} hours of exposure. "
                        "Error bars are Wilson score 95% confidence intervals.",
                width=Inches(3.5),
            )

        self._add_lc50_timeseries_section()

        timecourse_fig = self.analysis_results.get("timecourse_plot_figure")
        if timecourse_fig and self._wants("results_timecourse"):
            self._results_sub += 1
            self.document.add_heading(f"{self._results_no}.{self._results_sub} Mortality Time-Course", level=2)
            self.document.add_paragraph(
                f"Cumulative mortality was tracked across the daily observation timepoints ({self._timepoint_list()}) "
                "for each concentration group."
            )
            self._add_plot_to_document(
                timecourse_fig,
                caption="Cumulative mortality of Danio rerio embryos over the exposure period, by concentration group.",
                width=Inches(3.5),
            )

        fate_fig = self.analysis_results.get("fate_plot_figure")
        if fate_fig and self._wants("results_fate"):
            self._results_sub += 1
            self.document.add_heading(f"{self._results_no}.{self._results_sub} Fate Composition", level=2)
            self.document.add_paragraph(
                f"The developmental fate of embryos at {self._analysis_hpf()} hpf is summarized as the proportion of live embryos, "
                "live hatched larvae, dead embryos and dead hatched larvae per group."
            )
            self._add_plot_to_document(
                fate_fig,
                caption=f"Fate composition per treatment group at {self._analysis_hpf()} hours post-fertilization.",
                width=Inches(3.5),
            )

        self._add_sublethal_stats_section()

        hatching_fig = self.analysis_results.get("hatching_plot_figure")
        if hatching_fig and self._wants("results_hatching"):
            self._results_sub += 1
            self.document.add_heading(f"{self._results_no}.{self._results_sub} Hatching Rate", level=2)
            self.document.add_paragraph(
                f"The hatching rate of embryos was assessed at {self._analysis_hpf()} hours post-fertilization across all treatment groups."
            )
            self._add_plot_to_document(
                hatching_fig,
                caption=f"Hatching rate per treatment group at {self._analysis_hpf()} hours post-fertilization.",
                width=Inches(3.5),
            )

        malformation_fig = self.analysis_results.get("malformation_plot_figure")
        if malformation_fig and self._wants("results_malformation"):
            self._results_sub += 1
            self.document.add_heading(f"{self._results_no}.{self._results_sub} Malformation Profile", level=2)
            self.document.add_paragraph(
                "The incidence of specific morphological abnormalities was recorded to generate a malformation profile. A detailed breakdown of malformation frequencies is provided in the Appendices."
            )
            self._add_plot_to_document(
                malformation_fig,
                caption="Malformation profile detailing the incidence of sublethal effects per group.",
                width=Inches(5.0),
            )

    # Appendix Section #
    def _create_appendix_section(self) -> None:
        """Creates the appendix section with supplementary data."""
        if not any(self._wants(key) for key in
                   ("appendix_layout", "appendix_raw",
                    "appendix_malformation", "appendix_photos")):
            return
        self.document.add_page_break()
        appendices = [
            ("appendix_layout", "Experimental Plate Layout", self._add_plate_layout_appendix),
            ("appendix_raw", "Raw Observational Data", self._add_raw_data_appendix),
            ("appendix_malformation", "Malformation Frequency", self._add_malformation_frequency_appendix),
            ("appendix_photos", "Photographic Documentation", self._create_photo_appendix),
        ]
        selected = [(title, build) for key, title, build in appendices if self._wants(key)]
        if not selected:
            return

        self._begin_section("Appendices")
        # Lettered by position among those actually included, so omitting one
        # does not leave a gap in the sequence.
        for index, (title, build) in enumerate(selected):
            self.document.add_heading(f"Appendix {chr(65 + index)}: {title}", level=2)
            build()
        
    def _add_plate_layout_appendix(self) -> None:
        """Adds the plate layout map to the appendix."""
        plate_layouts = self.project_data.get("plate_layout", {})
        if not plate_layouts:
            self.document.add_paragraph("No plate layout data available.", style='Normal')
            return

        p_rows, p_cols = self.project_data.get("plate_dimensions", (8, 12))

        conc_map = self.project_data.get("concentration_map", {})
        conc_unit = self.project_data.get("concentration_unit", "")

        for plate_idx, layout in plate_layouts.items():
            self.table_count += 1
            p = self.document.add_paragraph(f"Table {self.table_count}. Well assignment map for Plate {int(plate_idx)}.")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)

            table = self.document.add_table(rows=p_rows + 1, cols=p_cols + 1)
            table.style = 'Table Grid'

            for j in range(1, p_cols + 1): table.cell(0, j).text = str(j)
            for i in range(1, p_rows + 1): table.cell(i, 0).text = chr(64 + i)

            for well_id, group_id in layout.items():
                row_char = well_id[0]
                col_num = int(well_id[1:])
                row_idx = ord(row_char) - 64
                conc = conc_map.get(group_id, {})
                value = conc.get("value", "")
                display = f"{group_id}\n{value} {conc_unit}".strip() if value != "" else group_id
                table.cell(row_idx, col_num).text = display

            self._style_table(table, font_size=Pt(8))

    def _add_raw_data_appendix(self) -> None:
        """Reference a sibling CSV holding the day-by-day observational data.

        A full FET project runs to thousands of rows — a six-plate, four-day study
        is 2160 — and emitting them as one Word table produces a document that is
        slow to open and impractical to read. The data is written alongside the
        report instead, where it can be opened in a spreadsheet or a statistics
        package. The table is kept as a fallback if that file cannot be written,
        so the observations are never dropped from the record.
        """
        well_data = self.project_data.get("well_data", {})
        if not well_data:
            self.document.add_paragraph("No raw observational data available.", style='Normal')
            return

        all_data = []
        conc_map = self.project_data.get("concentration_map", {})
        plate_layout = self.project_data.get("plate_layout", {})

        for day, plates in well_data.items():
            for plate_idx, wells in plates.items():
                for well_id, data in wells.items():
                    group_id = plate_layout.get(str(plate_idx), {}).get(well_id, 'N/A')
                    conc_info = conc_map.get(group_id, {})
                    row = {
                        "Day": int(day), "Plate": int(plate_idx), "Well": well_id,
                        "Group": group_id, "Conc.": conc_info.get('value', 'N/A'),
                        "Status": data.get("status", "Normal"),
                        "Hatched": "Yes" if data.get("status") in HATCHED_STATUSES else "No",
                        "Malformations": ", ".join(data.get("sublethal_conditions", [])),
                        "Notes": data.get("notes", "")
                    }
                    all_data.append(row)
        
        if not all_data:
            log.warning("No valid observational data found for raw data appendix.")
            return

        df = pd.DataFrame(all_data).sort_values(by=["Day", "Plate", "Well"])

        csv_name = self._write_raw_data_csv(df)
        if csv_name:
            self.document.add_paragraph(
                f"The complete observational record — {len(df)} well observations across "
                f"{df['Day'].nunique()} day(s) — accompanies this report as “{csv_name}”, "
                "saved in the same folder. Each row gives the day, plate, well, treatment "
                "group, nominal concentration, recorded status, hatching, morphological "
                "abnormalities and any note entered at the time of scoring.",
                style='Normal',
            )
            return

        self.table_count += 1
        p = self.document.add_paragraph(f"Table {self.table_count}. Raw observational data recorded daily for each well.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)

        table = self.document.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        for i, col_name in enumerate(df.columns): table.cell(0, i).text = col_name

        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row): row_cells[i].text = str(val)

        self._style_table(table, font_size=Pt(8))

    def _write_raw_data_csv(self, df: pd.DataFrame) -> Optional[str]:
        """Write *df* beside the report; return its filename, or None on failure."""
        if not self._output_path:
            return None
        stem = os.path.splitext(os.path.basename(self._output_path))[0]
        csv_name = f"{stem}_RawData.csv"
        csv_path = os.path.join(os.path.dirname(self._output_path) or ".", csv_name)
        try:
            df.to_csv(csv_path, index=False, encoding="utf-8-sig")
        except OSError as e:
            log.warning(f"Could not write the raw-data CSV beside the report: {e}")
            return None
        return csv_name

    def _add_malformation_frequency_appendix(self) -> None:
        """Frequency of each morphological endpoint per group, among survivors.

        Scored on the day the reported endpoints were computed for — not the last
        day of the project, which disagreed with this table's own caption whenever
        an earlier day was analysed.

        Only surviving embryos are counted, matching the analysis: a dead or
        coagulated embryo cannot be assessed for malformation, so including one
        here made this appendix contradict the summary table in the same report.
        The survivor count per group is given alongside so the frequencies can be
        read as proportions.
        """
        day = str(self._analysis_day())
        well_data_day = self.project_data.get("well_data", {}).get(day, {})

        if not well_data_day:
            self.document.add_paragraph(
                f"No malformation data available for Day {day}.", style='Normal'
            )
            return

        malformation_counts: Dict[str, Counter] = {}
        survivors: Counter = Counter()
        plate_layout = self.project_data.get("plate_layout", {})
        all_malformations = set()

        for plate_idx, wells in well_data_day.items():
            for well_id, data in wells.items():
                if data.get("status") not in LIVE_STATUSES:
                    continue
                group_id = plate_layout.get(str(plate_idx), {}).get(well_id, 'N/A')
                if group_id not in malformation_counts:
                    malformation_counts[group_id] = Counter()
                survivors[group_id] += 1

                malformations = data.get("sublethal_conditions", [])
                malformation_counts[group_id].update(malformations)
                all_malformations.update(malformations)

        if not all_malformations:
            self.document.add_paragraph(
                "No specific malformations were recorded among surviving embryos.",
                style='Normal',
            )
            return

        sorted_malformations = sorted(all_malformations)

        self.table_count += 1
        p = self.document.add_paragraph(
            f"Table {self.table_count}. Frequency of specific malformations among surviving "
            f"embryos at {self._analysis_hpf()} hpf."
        )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)

        headers = ["Group ID", "N (survivors)"] + sorted_malformations
        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        for i, header in enumerate(headers):
            table.cell(0, i).text = header

        for group_id in sorted(malformation_counts):
            counts = malformation_counts[group_id]
            row_cells = table.add_row().cells
            row_cells[0].text = str(group_id)
            row_cells[1].text = str(survivors[group_id])
            for i, malf in enumerate(sorted_malformations, start=2):
                row_cells[i].text = str(counts.get(malf, 0))

        self._style_table(table, font_size=Pt(9))
        
    def _create_photo_appendix(self) -> None:
        """Adds day-separated panels of representative photos to the appendix."""
        photos_with_metadata = self.project_data.get("photos_with_metadata", [])
        if not photos_with_metadata:
            self.document.add_paragraph("No photographic documentation was attached to this project.", style='Normal')
            return

        # Group photos by day
        photos_by_day: Dict[int, List[Dict[str, Any]]] = {}
        for meta in photos_with_metadata:
            day = int(meta.get("day", 0))
            photos_by_day.setdefault(day, []).append(meta)

        conc_map = self.project_data.get("concentration_map", {})
        plate_layouts = self.project_data.get("plate_layout", {})

        # Generate one panel per day
        for day, day_photos in sorted(photos_by_day.items()):
            panel_buffer, processed_photos = self._create_photo_panel(day_photos)
            if not panel_buffer:
                continue

            # Add the figure
            base_caption = f"Representative images of embryos recorded at Day {day}. "
            self._add_plot_to_document(panel_buffer, caption=base_caption, width=Inches(6.5))

            # Add the detailed, continuous legend
            legend_p = self.document.paragraphs[-1] # Get the caption paragraph just added
            for i, meta in enumerate(processed_photos):
                label = chr(65 + i)
                group_id = plate_layouts.get(str(meta['plate']), {}).get(meta['well'], "N/A")
                conc_info = conc_map.get(group_id, {})
                conc_value = conc_info.get('value', 'N/A')
                conc_unit = self.project_data.get("concentration_unit", "unit")

                run_label = legend_p.add_run(f"{label}: ")
                run_label.bold = True
                run_label.font.size = Pt(9)
                run_label.font.name = FONT_NAME

                well_info = (
                    self.project_data.get("well_data", {})
                    .get(str(meta["day"]), {})
                    .get(str(meta["plate"]), {})
                    .get(meta["well"], {})
                )
                status = well_info.get("status", "")
                conditions = well_info.get("sublethal_conditions", [])
                status_str = f"; Status: {status}" if status else ""
                cond_str = f"; Conditions: {', '.join(conditions)}" if conditions else ""

                legend_text = f"Plate {meta['plate']}, Well {meta['well']} ({group_id}, {conc_value} {conc_unit}){status_str}{cond_str}. "
                run_details = legend_p.add_run(legend_text)
                run_details.font.size = Pt(9)
                run_details.font.name = FONT_NAME
                
            # Final styling for the legend paragraph
            legend_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            legend_p.paragraph_format.space_before = Pt(0)
            legend_p.paragraph_format.space_after = Pt(12)


    # Tables #
    def _style_table(self, table: Table, font_size: Pt = Pt(FONT_SIZE_TABLE)) -> None:
        """Applies consistent styling (font, alignment, bold header) to a table."""
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = FONT_NAME
                        run.font.size = font_size
                        if i == 0:
                            run.bold = True

    def _add_key_value_table(self, title: str, data_dict: Dict[str, Any]) -> None:
        """Helper to create a simple two-column key-value table."""
        self.table_count += 1
        p = self.document.add_paragraph(f"Table {self.table_count}. {title}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)

        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text, hdr_cells[1].text = 'Parameter', 'Value'
        
        for key, value in data_dict.items():
            row_cells = table.add_row().cells
            row_cells[0].paragraphs[0].text = str(key)
            
            p_val = row_cells[1].paragraphs[0]
            run = p_val.add_run(str(value))
            if key == "Species":
                run.italic = True
        
        self._style_table(table)

    @staticmethod
    def _leading_float(text) -> Optional[float]:
        """Parse a leading number from a free-text measurement, else None."""
        import re
        if text is None:
            return None
        m = re.search(r"-?\d+(?:[.,]\d+)?", str(text))
        return float(m.group().replace(",", ".")) if m else None

    # OECD TG 236 in-test acceptability ranges, checked per parameter against the
    # monitoring log. The two labels differ because one heads an advisory
    # ("Dissolved-oxygen advisory:", hyphenated as a compound modifier) and the
    # other is a sentence subject ("Dissolved oxygen remained within ...").
    # Fields: (log key, subject, advisory heading, predicate, criterion, format).
    _WATER_QUALITY_CRITERIA = (
        ("temperature", "Temperature", "Temperature",
         lambda v: 25.0 <= v <= 27.0,
         "the OECD TG 236 range of 26 ± 1 °C", "{:g} °C"),
        ("dissolved_oxygen", "Dissolved oxygen", "Dissolved-oxygen",
         lambda v: v >= 80.0,
         "the OECD TG 236 minimum of 80% saturation", "{:g}%"),
    )

    def _water_quality_findings(self) -> List[Dict[str, Any]]:
        """Per-parameter outcome of the monitored water quality.

        Each parameter is assessed independently and only over the days on which
        it was actually recorded. A blank field is an unmeasured value, not a
        passing one, so a parameter with no parsable entry yields no finding at
        all — the report must never certify a validity criterion against data
        that was never collected.

        Structured rather than prose because both the advisory paragraphs and the
        validity verdict read it; one pass is what keeps them from disagreeing.
        """
        log_data = self.project_data.get("water_quality_log", {}) or {}
        if not log_data:
            return []

        days = sorted(log_data.keys(), key=lambda d: int(d))
        findings: List[Dict[str, Any]] = []
        for key, subject, advisory, in_range, criterion, value_format in self._WATER_QUALITY_CRITERIA:
            measured, out_of_range = [], []
            for day in days:
                value = self._leading_float((log_data[day] or {}).get(key))
                if value is None:
                    continue
                measured.append(day)
                if not in_range(value):
                    out_of_range.append(f"Day {day} ({value_format.format(value)})")

            if not measured:
                continue  # never recorded: make no claim either way
            findings.append({
                "key": key, "subject": subject, "advisory": advisory,
                "criterion": criterion, "measured": measured,
                "out_of_range": out_of_range, "passed": not out_of_range,
                "scope_is_every_day": len(measured) == len(days),
            })
        return findings

    def _water_quality_validity_notes(self) -> List[str]:
        """Advisory OECD TG 236 statements on the monitored water quality."""
        notes: List[str] = []
        for f in self._water_quality_findings():
            if f["out_of_range"]:
                notes.append(
                    f"{f['advisory']} advisory: the following day(s) fell outside "
                    f"{f['criterion']}: " + ", ".join(f["out_of_range"]) + "."
                )
            else:
                scope = ("all recorded days" if f["scope_is_every_day"]
                         else "day(s) " + ", ".join(str(d) for d in f["measured"]))
                notes.append(f"{f['subject']} remained within {f['criterion']} on {scope}.")
        return notes

    def _internal_plate_control_notes(self) -> List[str]:
        """Advisory statements on plates failing the TG 236 §23 internal control.

        The analysis reports such plates but does not exclude them: rejecting a
        plate withdraws whichever concentrations it carried from the LC50, which is
        the operator's decision. The note therefore states the observation and the
        guideline's consequence, and leaves the verdict alone.
        """
        failures = self.analysis_results.get("plate_control_failures") or []
        if not failures:
            return []
        plates = ", ".join(
            f"plate {f['plate']} ({f['dead']} of {f['n']} wells)" for f in failures
        )
        subject = "This plate" if len(failures) == 1 else "These plates"
        return [
            f"Internal plate control advisory: more than one dead embryo was recorded in the "
            f"internal (dilution-water) control of {plates}. OECD TG 236 §23 indicates that "
            f"{subject.lower()} should be rejected, which would reduce the concentrations "
            f"available for the LC50. The endpoints reported here include all plates."
        ]

    def _add_water_quality_table(self) -> None:
        """Per-day physicochemical monitoring table (OECD TG 236 in-test monitoring)."""
        log_data = self.project_data.get("water_quality_log", {}) or {}
        if not log_data:
            return

        self.table_count += 1
        p = self.document.add_paragraph(
            f"Table {self.table_count}. Physicochemical water-quality parameters monitored during the test."
        )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)

        headers = ["Day", "Temperature (°C)", "Dissolved O₂", "pH", "Conductivity (µS/cm)", "Notes"]
        keys = ["temperature", "dissolved_oxygen", "ph", "conductivity", "notes"]
        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.cell(0, i).text = h
        for day in sorted(log_data.keys(), key=lambda d: int(d)):
            entry = log_data[day]
            cells = table.add_row().cells
            cells[0].text = str(day)
            for i, k in enumerate(keys, start=1):
                cells[i].text = str(entry.get(k, "") or "")
        self._style_table(table, font_size=Pt(9))

    def _add_results_summary_table(self, summary_df: pd.DataFrame) -> None:
        """Creates the main results summary table."""
        if summary_df is None or summary_df.empty:
            self.document.add_paragraph(
                "Results summary data is unavailable. Ensure all experiment days have been "
                "recorded and analysis has been run before generating the report."
            )
            return

        self.table_count += 1
        p = self.document.add_paragraph(
            f"Table {self.table_count}. Summary of lethal and sublethal effects at {self._analysis_hpf()} hours post-fertilization."
        )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        
        report_df = summary_df.copy()
        # Mortality and hatching are expressed per embryo actually scored, not per
        # well assigned; see biostatistics.evaluable_n.
        report_df["n_scored"] = evaluable_n(report_df)
        scored_denom = report_df["n_scored"].replace(0, float("nan"))
        report_df["Mortality (%)"] = (report_df["dead"] / scored_denom * 100).fillna(0)
        report_df["Hatched (%)"] = (report_df["hatched"] / scored_denom * 100).fillna(0)
        # Malformation is expressed as a percentage of surviving (live) embryos.
        malf_denom = report_df["live"] if "live" in report_df.columns else report_df["total"]
        report_df["Malformed (%)"] = (
            report_df["malformed"] / malf_denom.replace(0, float("nan")) * 100
        ).fillna(0)

        has_abbott = "mortality_abbott" in report_df.columns
        if has_abbott:
            # Keep NaN (e.g. the Positive Control row, which is not background-corrected)
            # so it renders blank rather than 0.
            report_df["Abbott-corrected Mortality (%)"] = report_df["mortality_abbott"]

        conc_unit = self.project_data.get("concentration_unit", "unit")
        report_df = report_df.rename(columns={"conc_id": "Group ID"})
        cols_to_show = ["Group ID", "conc_value", "total", "n_scored", "dead", "Mortality (%)"]
        if has_abbott:
            cols_to_show.append("Abbott-corrected Mortality (%)")
        cols_to_show += ["Hatched (%)", "Malformed (%)"]
        report_df = report_df[cols_to_show]


        table = self.document.add_table(rows=1, cols=len(cols_to_show))
        table.style = 'Table Grid'
        header_map = {"conc_value": f"Conc. ({conc_unit})", "total": "N assigned",
                      "n_scored": "N scored", "dead": "Dead"}
        for i, col_name in enumerate(report_df.columns):
            table.cell(0, i).text = header_map.get(col_name, col_name)

        for _, row in report_df.iterrows():
            row_cells = table.add_row().cells
            for i, (col_name, val) in enumerate(row.items()):
                if isinstance(val, float) and '%' in report_df.columns[i]:
                    row_cells[i].text = "" if pd.isna(val) else f"{val:.2f}"
                elif col_name == 'conc_value':
                    row_cells[i].text = f"{val:.4f}".rstrip('0').rstrip('.')
                elif isinstance(val, (int, float)):
                    row_cells[i].text = str(int(val))
                else:
                    row_cells[i].text = str(val)
                    
        self._style_table(table)


    # Figures #
    def _add_plot_to_document(self, figure: Union[plt.Figure, io.BytesIO], caption: str, width=Inches(6.0)) -> None:
        """Adds a matplotlib figure to the document with a caption."""
        self.figure_count += 1
        mem_fig = io.BytesIO()
        if isinstance(figure, plt.Figure):
            figure.savefig(mem_fig, format='png', dpi=600, bbox_inches='tight')
            plt.close(figure) # Close the figure to free up memory
            mem_fig.seek(0)
        else:
            mem_fig.seek(0) # Ensure buffer is at the start
            mem_fig = figure

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(mem_fig, width=width)
        
        p_caption = self.document.add_paragraph() # Empty paragraph for the caption
        p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_caption.paragraph_format.space_before = Pt(6)
        p_caption.paragraph_format.space_after = Pt(12)
        
        # Add "Figure X." in bold, then the rest of the caption
        run_fig_num = p_caption.add_run(f"Figure {self.figure_count}. ")
        run_fig_num.bold = True
        run_fig_num.font.size = Pt(10)
        run_fig_num.font.name = FONT_NAME
        
        run_caption = p_caption.add_run(caption)
        run_caption.font.size = Pt(10)
        run_caption.font.name = FONT_NAME
        
    def _noec_correction_label(self) -> str:
        """Name the multiplicity correction the NOEC/LOEC was derived under.

        Holm and Bonferroni can disagree about the LOEC, so the reader has to be
        told which produced the reported value.
        """
        label = self.analysis_results.get("noec_loec_results", {}).get("correction_label")
        if label == "Holm-Bonferroni":
            self._cite("holm")
            return label
        self._cite("dunn")
        return label or "Bonferroni"

    def _tsk_methods_sentence(self) -> str:
        """Describe the non-parametric estimate, when one was obtained."""
        if self.analysis_results.get("tsk_results", {}).get("lc50_numeric") is None:
            return ""
        self._cite("hamilton")
        trim = self.analysis_results["tsk_results"].get("trim")
        trim_clause = (
            f"with the smallest trim the observed response supported ({trim:.0%})"
            if trim is not None else "with an automatically selected trim"
        )
        return (
            f" A non-parametric LC50 was additionally estimated by the trimmed "
            f"Spearman-Karber method (Hamilton et al., 1977), {trim_clause}; it "
            "requires no distributional assumption and yields an estimate where "
            "curve fitting does not converge."
        )

    def _validity_criteria_sentence(self) -> str:
        """Name the TG 236 §9 criteria this analysis is in a position to assess.

        Naming only control mortality described one criterion of six as though it
        were the guideline's requirement, in the section a reader consults to
        understand the method.
        """
        return (
            "Test validity was assessed against the OECD TG 236 criteria that could be "
            "evaluated from the recorded data: control survival, control hatching rate at "
            "the end of the exposure, positive-control sensitivity where a positive control "
            "was included, the fertilization rate of the egg batch where recorded, and the "
            "monitored temperature and dissolved oxygen. "
        )

    def _species(self) -> str:
        """Scientific name of the test species."""
        return (self.project_data.get("test_organisms", {}) or {}).get("species") or DEFAULT_SPECIES

    def _add_mortality_bounds_sentence(self) -> None:
        """Report the concentrations bracketing the observed lethal range (TG 236 §42).

        Stated only for the bounds the design actually produced: a series where no
        group reached 0% or 100% mortality has no such concentration to report.
        """
        bounds = mortality_bounding_concentrations(
            self.analysis_results.get("summary_df", pd.DataFrame())
        )
        conc_unit = self.project_data.get("concentration_unit", "unit")
        hpf = self._analysis_hpf()
        clauses = []
        if bounds["no_mortality_max"] is not None:
            clauses.append(
                f"the highest concentration causing no mortality was "
                f"{bounds['no_mortality_max']:.4f} {conc_unit}"
            )
        if bounds["full_mortality_min"] is not None:
            clauses.append(
                f"the lowest concentration causing 100% mortality was "
                f"{bounds['full_mortality_min']:.4f} {conc_unit}"
            )
        if clauses:
            self.document.add_paragraph(
                f"At {hpf} hours of exposure, " + " and ".join(clauses) + "."
            )

    def _add_tsk_sentence(self) -> None:
        """Report the non-parametric LC50 alongside the fitted one."""
        tsk = self.analysis_results.get("tsk_results", {})
        if tsk.get("lc50_numeric") is None:
            return
        conc_unit = self.project_data.get("concentration_unit", "unit")
        fitted = bool(self.analysis_results.get("lc50_results", {}).get("_fitted_params"))
        if fitted:
            self.document.add_paragraph(
                f"The trimmed Spearman-Karber estimate of the LC50 was {tsk['lc50']} "
                f"{conc_unit}, provided as a distribution-free check on the fitted value."
            )
        else:
            self.document.add_paragraph(
                f"As the logistic model did not converge, the LC50 is reported from the "
                f"trimmed Spearman-Karber estimate: {tsk['lc50']} {conc_unit}."
            )

    def _weighting_sentence(self) -> str:
        """State how the groups were weighted in the fit.

        Group sizes differ across a FET design, so whether the fit accounted for
        that determines the LC50 a reader would reproduce.
        """
        if self.analysis_results.get("lc50_results", {}).get("weighting") != "binomial":
            return ""
        return (
            " Groups were weighted by the inverse of their binomial variance, so that a "
            "concentration scored in few surviving embryos carries proportionally less "
            "influence on the fitted curve than one scored in many."
        )

    def _build_lc50_methods_sentence(self, model_info: dict, bootstrap_method: str = None) -> str:
        self._cite("efron", "ritz")
        if bootstrap_method == "rigorous":
            _ci_note = (
                "Confidence intervals were estimated by a parametric (binomial) bootstrap (Efron, 1979; "
                "500 iterations, seed = 15) that resampled control and treatment mortality and re-applied "
                "Abbott's correction at each iteration, propagating control-mortality uncertainty into the "
                "interval."
            )
        else:
            _ci_note = ("Confidence intervals were estimated by case-resampling bootstrap "
                        "(Efron, 1979; 500 iterations, seed = 15).")

        if not model_info:
            return ("LC50 values were estimated by fitting mortality data to a four-parameter logistic (4PL) "
                    f"model using non-linear regression (Ritz, 2010). {_ci_note}")

        mode = model_info.get("mode", "manual")
        display = model_info.get("display_name", "4PL (all free)")
        bottom = model_info.get("bottom")
        top = model_info.get("top")
        n_free = model_info.get("n_free", 4)

        if mode == "auto":
            self._cite("burnham_anderson")
            return (
                f"To estimate the LC50, two- (2PL), three- (3PL), and four-parameter (4PL) logistic models were fitted to "
                f"the mortality data by non-linear regression (Ritz, 2010). Model selection was based on the corrected Akaike Information Criterion (AICc), which "
                f"penalizes model complexity relative to the small number of concentration groups typical of FET assays "
                f"(Burnham & Anderson, 2002). The best-fitting model was {display}. {_ci_note}"
            )
        if n_free == 4:
            return f"LC50 values were estimated by fitting mortality data to a four-parameter logistic (4PL) model using non-linear regression (Ritz, 2010). {_ci_note}"
        if n_free == 2:
            return (
                f"LC50 values were estimated by fitting mortality data to a two-parameter logistic (2PL) model using "
                f"non-linear regression (Ritz, 2010), with the bottom asymptote constrained at {bottom:.1f}% and the top asymptote "
                f"constrained at {top:.1f}%. {_ci_note}"
            )
        if bottom is not None:
            return (
                f"LC50 values were estimated by fitting mortality data to a three-parameter logistic (3PL) model using "
                f"non-linear regression (Ritz, 2010), with the bottom asymptote constrained at {bottom:.1f}%. {_ci_note}"
            )
        if top is not None:
            return (
                f"LC50 values were estimated by fitting mortality data to a three-parameter logistic (3PL) model using "
                f"non-linear regression (Ritz, 2010), with the top asymptote constrained at {top:.1f}%. {_ci_note}"
            )
        return f"LC50 values were estimated by fitting mortality data to a logistic model using non-linear regression (Ritz, 2010). {_ci_note}"

    def _add_trend_test_sentence(self) -> None:
        trend_results = self.analysis_results.get("trend_results", {})
        trend = trend_results.get("trend", "Not Calculated")
        p_value = trend_results.get("p_value", "Not Calculated")
        statistic = trend_results.get("statistic", "Not Calculated")
        if p_value not in ("Not Calculated", None):
            self.document.add_paragraph(
                f"A Cochran-Armitage trend test yielded a test statistic of Z = {statistic} "
                f"(one-sided p = {p_value}). {trend}"
            )
        elif trend not in ("Not Calculated", None):
            self.document.add_paragraph(
                "The Cochran-Armitage trend test could not be computed: "
                + (trend if str(trend).endswith(".") else f"{trend}.")
            )

    def _add_teratogenic_index_sentence(self) -> None:
        conc_unit = self.project_data.get("concentration_unit", "unit")
        ti = self.analysis_results.get("teratogenic_index", {})
        ti_val = ti.get("teratogenic_index", "Not Calculated")
        ec50 = ti.get("ec50_malformation", "Not Calculated")
        if ti.get("ti_numeric") is not None:
            self.document.add_paragraph(
                f"The median effect concentration for malformation (EC50) was {ec50} {conc_unit}, giving a "
                f"teratogenic index (TI = LC50/EC50) of {ti_val}. {ti.get('interpretation', '')}"
            )

        pooled = self.analysis_results.get("sublethal_stats", {}).get("pooled", {})
        loec = pooled.get("loec")
        noec = pooled.get("noec")
        if pooled.get("no_events"):
            # No abnormality anywhere, controls included. Reporting a NOEC at the
            # top concentration would imply an effect had been sought and bounded.
            self.document.add_paragraph(
                "No morphological abnormalities were recorded in any group, in the controls or at any test "
                "concentration; a sublethal NOEC/LOEC is therefore not defined."
            )
        elif loec and loec not in ("Not Calculated", None):
            noec_disp = format_with_unit(pooled.get("noec_numeric"), noec, conc_unit)
            loec_disp = format_with_unit(pooled.get("loec_numeric"), loec, conc_unit)
            self.document.add_paragraph(
                "For the pooled sublethal endpoint (embryos presenting at least one abnormality), the sublethal "
                f"NOEC and LOEC were {noec_disp} and {loec_disp}, respectively (Fisher's Exact Test, "
                "Benjamini-Hochberg corrected)."
            )

    def _add_mortality_effects_table(self, summary_df: pd.DataFrame) -> None:
        """Per-group mortality with Wilson 95% CI and the odds ratio (95% CI) vs
        the pooled control."""
        if summary_df is None or summary_df.empty:
            return
        controls = select_control_rows(summary_df, self._control_mode())
        if controls.empty:
            return
        ctrl_dead = int(controls["dead"].sum())
        ctrl_total = int(evaluable_n(controls).sum())
        if ctrl_total == 0:
            return

        conc_unit = self.project_data.get("concentration_unit", "unit")
        rows = summary_df[summary_df["conc_type"] == "Substrate"].sort_values("conc_value")
        if rows.empty:
            return

        self.table_count += 1
        p = self.document.add_paragraph(
            f"Table {self.table_count}. Mortality effect sizes at {self._analysis_hpf()} hpf: Wilson score 95% CIs and the "
            "odds ratio versus the pooled control."
        )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)

        headers = [f"Conc. ({conc_unit})", "Mortality % (95% CI)", "Odds ratio (95% CI)"]
        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.cell(0, i).text = h

        for _, r in rows.iterrows():
            k, n = int(r["dead"]), int(evaluable_n(r))
            c, lo, hi = wilson_ci(k, n)
            or_, or_lo, or_hi = odds_ratio_ci(k, n, ctrl_dead, ctrl_total)
            conc_txt = f"{r['conc_value']:.4f}".rstrip("0").rstrip(".")
            cells = table.add_row().cells
            cells[0].text = conc_txt
            cells[1].text = f"{c:.1f} ({lo:.1f}–{hi:.1f})"
            cells[2].text = (f"{or_:.2f} ({or_lo:.2f}–{or_hi:.2f})"
                             if or_ == or_ else "—")  # nan check
        self._style_table(table)

    def _add_lc50_timeseries_section(self) -> None:
        """LC50 at each daily timepoint, when the series was computed.

        A single-day LC50 characterizes potency at that timepoint; the series
        shows how it moves with exposure duration, which is how acute fish
        toxicity is conventionally reported.
        """
        series = self.analysis_results.get("lc50_timeseries")
        if not series or not self._wants("results_timeseries"):
            return

        conc_unit = self.project_data.get("concentration_unit", "unit")
        self._results_sub += 1
        self.document.add_heading(f"{self._results_no}.{self._results_sub} LC50 Time-Series", level=2)
        self.document.add_paragraph(
            "The LC50 was estimated independently at each daily observation timepoint, "
            "using the same model, asymptote constraints, background correction and "
            "reference control as the primary analysis, so the values are directly "
            "comparable with one another."
        )

        self.table_count += 1
        p = self.document.add_paragraph(
            f"Table {self.table_count}. LC50 by exposure time."
        )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)

        headers = ["Time (hpf)", f"LC50 ({conc_unit})", "95% CI", "Model",
                   f"Spearman-Kärber ({conc_unit})", "Dose groups"]
        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, header in enumerate(headers):
            table.cell(0, i).text = header

        for entry in series:
            cells = table.add_row().cells
            cells[0].text = str(entry["hpf"])
            if entry["lc50_numeric"] is None:
                # Naming the reason keeps a timepoint the data could not support
                # from reading as one where the compound had no effect.
                cells[1].text = "not estimable"
                cells[2].text = "—"
            else:
                cells[1].text = f"{entry['lc50_numeric']:.4f}"
                cells[2].text = format_ci_range(entry.get("ci_low"), entry.get("ci_high"))
            cells[3].text = str(entry.get("model") or "—")
            tsk = entry.get("tsk_numeric")
            cells[4].text = f"{tsk:.4f}" if tsk is not None else "—"
            cells[5].text = str(entry.get("n_groups", ""))

        self._style_table(table, font_size=Pt(9))

        figure = self.analysis_results.get("lc50_timeseries_figure")
        if figure is not None:
            self._add_plot_to_document(
                figure,
                caption="LC50 of Danio rerio embryos as a function of exposure time. "
                        "Error bars are bootstrap 95% confidence intervals.",
                width=Inches(3.5),
            )

    def _add_sublethal_stats_section(self) -> None:
        stats = self.analysis_results.get("sublethal_stats", {})
        # Gated on the per-endpoint tests rather than on 'available': with no
        # morphological endpoint scored, 'available' is still true through the
        # pooled endpoint and this section would be a heading plus a methods
        # paragraph describing tests that were never run. The pooled result is
        # reported under 2.2 either way.
        if not stats or not stats.get("tests") or not self._wants("results_sublethal"):
            return

        self._results_sub += 1
        self.document.add_heading(f"{self._results_no}.{self._results_sub} Sublethal Endpoint Analysis", level=2)
        self.document.add_paragraph(
            "Each morphological endpoint was compared with the control by Fisher's Exact Test (one-sided) with "
            "Benjamini-Hochberg control of the false discovery rate; odds ratios with 95% confidence intervals "
            "quantify the effect size. Significance is flagged as *** (p < 0.001), ** (p < 0.01), * (p < 0.05) "
            "or ns."
        )

        tests = [t for t in stats.get("tests", []) if t.get("p_adj") is not None]
        if tests:
            self.table_count += 1
            p = self.document.add_paragraph(
                f"Table {self.table_count}. Per-endpoint Fisher's Exact Test (Benjamini-Hochberg corrected) "
                f"versus control at {self._analysis_hpf()} hpf."
            )
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            headers = ["Endpoint", "Group", "n/N", "p (BH)", "Sig.", "OR (95% CI)"]
            table = self.document.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            for i, h in enumerate(headers):
                table.cell(0, i).text = h
            ca = stats.get("ca", {})
            for t in tests:
                cells = table.add_row().cells
                cells[0].text = str(t["endpoint"])
                cells[1].text = str(t["conc_id"])
                cells[2].text = f"{t['k']}/{t['n']}"
                cells[3].text = f"{t['p_adj']:.3g}"
                cells[4].text = significance_marker(t["p_adj"])
                cells[5].text = (f"{t['or']:.2f} ({t['or_lo']:.2f}–{t['or_hi']:.2f})"
                                 if t.get("or") == t.get("or") and t.get("or") is not None else "—")
            self._style_table(table, font_size=Pt(9))

            # Per-endpoint trend summary
            if ca:
                trend_bits = [
                    f"{ep} (Z = {r['z']:.2f}, p = {r['p']:.3g}, {significance_marker(r['p'])})"
                    for ep, r in ca.items()
                ]
                self.document.add_paragraph(
                    "Cochran-Armitage trend tests per endpoint: " + "; ".join(trend_bits) + "."
                )

    def _add_curve_fitting_section(self) -> None:
        lc50_results = self.analysis_results.get("lc50_results", {})
        model_info = lc50_results.get("model_info")
        lc50_val = lc50_results.get("lc50", "")
        lc50_is_numeric = bool(lc50_results.get("_fitted_params"))

        if not model_info or not lc50_is_numeric:
            return

        self.document.add_heading("Curve-Fitting Details", level=3)

        bottom = model_info.get("bottom")
        top = model_info.get("top")
        fitted = lc50_results.get("_fitted_params")

        details = {
            "Model": model_info.get("display_name", "—"),
            "Bottom asymptote": f"{bottom:.1f}% (fixed)" if bottom is not None else f"{fitted[0]:.2f}% (fitted)" if fitted else "fitted",
            "Top asymptote": f"{top:.1f}% (fixed)" if top is not None else f"{fitted[1]:.2f}% (fitted)" if fitted else "fitted",
            "Slope": lc50_results.get("slope", "—"),
            "EC50 / LC50": lc50_val,
            "Goodness of fit (R²)": lc50_results.get("r_squared", "—"),
        }
        self._add_key_value_table("Summary of dose-response curve fitting.", details)

        if model_info.get("mode") == "auto" and model_info.get("aic_table"):
            self._add_model_comparison_table(model_info["aic_table"])

    def _add_model_comparison_table(self, aic_table: list) -> None:
        self.table_count += 1
        p = self.document.add_paragraph(
            f"Table {self.table_count}. Model comparison by corrected Akaike Information Criterion (AICc). "
            f"The selected model (ΔAICc = 0.00) is shown in bold."
        )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)

        table = self.document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["Model", "Free parameters (k)", "AICc", "ΔAICc"]
        for i, h in enumerate(headers):
            table.cell(0, i).text = h

        best_row_idx = next((i for i, e in enumerate(aic_table) if e.get("delta", 1) == 0.0), None)

        for row_idx, entry in enumerate(aic_table):
            row_cells = table.add_row().cells
            row_cells[0].text = entry.get("model", "")
            row_cells[1].text = str(entry.get("k", ""))
            aicc_val = entry.get("aicc", float("inf"))
            # Naming the reason keeps a model that the design cannot support from
            # reading as one that simply scored poorly.
            if entry.get("estimable", aicc_val != float("inf")):
                row_cells[2].text = f"{aicc_val:.2f}"
                row_cells[3].text = f"{entry.get('delta', 0.0):.2f}"
            else:
                row_cells[2].text = "not estimable (n ≤ k+1)"
                row_cells[3].text = "—"
            if row_idx == best_row_idx:
                for cell in row_cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

        self._style_table(table, font_size=Pt(FONT_SIZE_TABLE))

    # Photo Panel #
    def _create_photo_panel(self, photos_meta: List[Dict[str, Any]], max_images: int = 12, columns: int = 3
                        ) -> Tuple[Optional[io.BytesIO], List[Dict[str, Any]]]:
        """Creates a composite image panel from individual well photos, maintaining aspect ratio."""
        paths_to_process = photos_meta[:max_images]
        if not paths_to_process:
            return None, []

        processed_images = []
        processed_meta: List[Dict[str, Any]] = []
        temp_files = []
        thumb_width = 400
        font = None
        font_candidates = [
            resource_path("resources/fonts/Inter/static/Inter_18pt-Bold.ttf"),
            "arialbd.ttf",
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
        ]
        for font_path in font_candidates:
            try:
                font = ImageFont.truetype(font_path, 50)
                break
            except (IOError, OSError):
                continue
        if font is None:
            log.warning("No suitable bold font found. Using default font.")
            font = ImageFont.load_default(size=50)

        # 1. Load, resize, and label all images first. Labels are assigned by
        #    successful-load order (not raw index) so an image that fails to
        #    load leaves no lettering gap and no orphan legend entry — the panel
        #    letters stay in lock-step with the returned metadata below.
        for meta in paths_to_process:
            try:
                _parts = meta['path'].replace("\\", "/").split("/")
                img_path = os.path.join(self.project_dir, *_parts)
                
                # Handle TIFF conversion
                ext = os.path.splitext(img_path)[1].lower()
                if ext in ['.tif', '.tiff']:
                    with Image.open(img_path) as im:
                        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                            im.save(tmp.name, "PNG")
                            img_path_to_open = tmp.name
                            temp_files.append(tmp.name)
                else:
                    img_path_to_open = img_path
                
                with Image.open(img_path_to_open) as img:
                    img = img.convert("RGBA")
                    
                    # Resize proportionally to fixed width
                    w_percent = (thumb_width / float(img.size[0]))
                    h_size = int((float(img.size[1]) * float(w_percent)))
                    img = img.resize((thumb_width, h_size), Image.Resampling.LANCZOS)

                    # Draw label directly on the resized image
                    draw = ImageDraw.Draw(img)
                    label = chr(65 + len(processed_images))
                    padding = 15
                    x, y = padding, padding
                    
                    outline_color = (255, 255, 255)
                    fill_color = (0, 0, 0)
                    
                    # Create a 2px outline for better contrast
                    for offset in [(-2,-2), (-2,2), (2,-2), (2,2)]:
                        draw.text((x+offset[0], y+offset[1]), label, font=font, fill=outline_color)
                    draw.text((x, y), label, font=font, fill=fill_color)

                    processed_images.append(img)
                    processed_meta.append(meta)

            except Exception as e:
                log.warning(f"Could not process image {meta['path']}: {e}")
                continue
        
        # Cleanup temporary files
        for f in temp_files:
            try: os.remove(f)
            except OSError as e: log.warning(f"Could not remove temp file {f}: {e}")

        if not processed_images:
            return None, []

        # 2. Calculate panel dimensions
        rows = (len(processed_images) + columns - 1) // columns
        panel_width = columns * thumb_width
        
        row_heights = []
        for r in range(rows):
            images_in_row = processed_images[r*columns : (r+1)*columns]
            if images_in_row:
                max_h = max(img.height for img in images_in_row)
                row_heights.append(max_h)
        
        panel_height = sum(row_heights)

        # 3. Create panel and paste images
        panel = Image.new('RGBA', (panel_width, panel_height), (255, 255, 255, 255))
        
        current_y = 0
        for r in range(rows):
            images_in_row = processed_images[r*columns : (r+1)*columns]
            current_x = 0
            for img in images_in_row:
                panel.paste(img, (current_x, current_y))
                current_x += img.width
            if images_in_row:
                current_y += row_heights[r]

        # 4. Save to buffer and return
        buffer = io.BytesIO()
        panel.save(buffer, format='PNG')
        buffer.seek(0)
        return buffer, processed_meta