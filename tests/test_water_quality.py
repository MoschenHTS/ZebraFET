"""
test_water_quality.py — Per-day water-quality monitoring log (OECD TG 236).

Covers persistence via ProjectManager, inclusion in the report snapshot, and the
report's monitoring table plus the advisory DO/temperature validity notes.
"""
import os

import pytest
from docx import Document

from src.core.project_manager import ProjectManager
from src.export.report_generator import ReportGenerator


INITIAL_DATA = {"project_name": "WQTest", "num_days": 4, "num_plates": 1}


@pytest.fixture
def manager(tmp_path):
    m = ProjectManager.create_new(str(tmp_path / "WQTest"), INITIAL_DATA)
    yield m
    m.close()


class TestWaterQualityPersistence:
    def test_empty_log_by_default(self, manager):
        assert manager.get_water_quality_log() == {}

    def test_save_and_get_round_trip(self, manager):
        manager.save_water_quality(1, temperature="26.1", dissolved_oxygen="92",
                                   ph="7.4", conductivity="520", notes="ok")
        log = manager.get_water_quality_log()
        assert log[1]["temperature"] == "26.1"
        assert log[1]["dissolved_oxygen"] == "92"
        assert log[1]["ph"] == "7.4"

    def test_upsert_overwrites_same_day(self, manager):
        manager.save_water_quality(2, temperature="25")
        manager.save_water_quality(2, temperature="27", ph="7.0")
        log = manager.get_water_quality_log()
        assert log[2]["temperature"] == "27"
        assert log[2]["ph"] == "7.0"

    def test_persists_across_reopen(self, tmp_path):
        proj = str(tmp_path / "WQPersist")
        m1 = ProjectManager.create_new(proj, {"project_name": "WQPersist"})
        m1.save_water_quality(1, temperature="26", dissolved_oxygen="88")
        m1.close()
        m2 = ProjectManager(proj)
        log = m2.get_water_quality_log()
        m2.close()
        assert log[1]["dissolved_oxygen"] == "88"

    def test_snapshot_includes_log(self, manager):
        manager.save_water_quality(1, temperature="26")
        snap = manager.get_report_snapshot()
        assert "water_quality_log" in snap
        assert snap["water_quality_log"][1]["temperature"] == "26"


def _snapshot_with_log(tmp_path, log_entries):
    m = ProjectManager.create_new(str(tmp_path / "WQReport"), {"project_name": "WQReport", "num_days": 4})
    for day, kwargs in log_entries.items():
        m.save_water_quality(day, **kwargs)
    snap = m.get_report_snapshot()
    m.close()
    return snap


class TestWaterQualityReport:
    def _make_report(self, tmp_path, snap):
        analysis = {"summary_df": __import__("pandas").DataFrame(),
                    "lc50_results": {"lc50": "Not Calculated",
                                     "model_info": {"display_name": "4PL", "mode": "manual", "n_free": 4}},
                    "noec_loec_results": {"noec": "NC", "loec": "NC"}}
        out = str(tmp_path / "wq_report.docx")
        ReportGenerator(snap, str(tmp_path), analysis).generate_report(out)
        return Document(out)

    def test_report_has_monitoring_table(self, tmp_path):
        snap = _snapshot_with_log(tmp_path, {1: {"temperature": "26", "dissolved_oxygen": "92"}})
        doc = self._make_report(tmp_path, snap)
        text = " ".join(p.text for p in doc.paragraphs)
        assert "water-quality parameters" in text.lower()

    def test_out_of_range_temperature_flagged(self, tmp_path):
        snap = _snapshot_with_log(tmp_path, {1: {"temperature": "30", "dissolved_oxygen": "92"}})
        doc = self._make_report(tmp_path, snap)
        text = " ".join(p.text for p in doc.paragraphs)
        assert "Temperature advisory" in text

    def test_low_do_flagged(self, tmp_path):
        snap = _snapshot_with_log(tmp_path, {1: {"temperature": "26", "dissolved_oxygen": "70"}})
        doc = self._make_report(tmp_path, snap)
        text = " ".join(p.text for p in doc.paragraphs)
        assert "Dissolved-oxygen advisory" in text

    def test_in_range_reports_compliance(self, tmp_path):
        snap = _snapshot_with_log(tmp_path, {1: {"temperature": "26", "dissolved_oxygen": "95"},
                                             2: {"temperature": "25.5", "dissolved_oxygen": "88"}})
        doc = self._make_report(tmp_path, snap)
        text = " ".join(p.text for p in doc.paragraphs)
        assert "Temperature remained within the OECD TG 236 range of 26 ± 1 °C" in text
        assert "Dissolved oxygen remained within the OECD TG 236 minimum of 80% saturation" in text
        assert "on all recorded days" in text

    def test_no_log_no_notes(self, tmp_path):
        snap = _snapshot_with_log(tmp_path, {})
        doc = self._make_report(tmp_path, snap)
        text = " ".join(p.text for p in doc.paragraphs)
        assert "advisory" not in text.lower()
        assert "water-quality parameters" not in text.lower()


class TestWaterQualityNeverCertifiesUnmeasuredData:
    """A blank field is an unmeasured value, not a passing one.

    The report is a citable artifact, so it must not state that an OECD TG 236
    acceptability criterion was met on the strength of a field nobody filled in.
    Each parameter is judged independently and only over the days it was recorded.
    """

    def _report_text(self, tmp_path, log_entries):
        snap = _snapshot_with_log(tmp_path, log_entries)
        analysis = {"summary_df": __import__("pandas").DataFrame(),
                    "lc50_results": {"lc50": "Not Calculated",
                                     "model_info": {"display_name": "4PL", "mode": "manual", "n_free": 4}},
                    "noec_loec_results": {"noec": "NC", "loec": "NC"}}
        out = str(tmp_path / "wq_claims.docx")
        ReportGenerator(snap, str(tmp_path), analysis).generate_report(out)
        return " ".join(p.text for p in Document(out).paragraphs)

    def test_blank_row_makes_no_compliance_claim(self, tmp_path):
        """Opening the dialog and saving it untouched must not certify anything."""
        text = self._report_text(tmp_path, {1: {}})
        assert "remained within" not in text

    def test_ph_only_makes_no_temperature_or_do_claim(self, tmp_path):
        text = self._report_text(tmp_path, {1: {"ph": "7.4"}})
        assert "Temperature remained within" not in text
        assert "Dissolved oxygen remained within" not in text

    def test_notes_only_makes_no_claim(self, tmp_path):
        text = self._report_text(tmp_path, {1: {"notes": "meter unavailable"}})
        assert "remained within" not in text

    def test_parameters_are_judged_independently(self, tmp_path):
        """Temperature recorded, dissolved oxygen never: claim one, not the other."""
        text = self._report_text(tmp_path, {1: {"temperature": "26"}, 2: {"temperature": "26.2"}})
        assert "Temperature remained within" in text
        assert "Dissolved oxygen remained within" not in text

    def test_compliance_scope_names_only_the_measured_days(self, tmp_path):
        """Day 2 carries no temperature, so the claim must not cover it."""
        text = self._report_text(tmp_path, {1: {"temperature": "26", "dissolved_oxygen": "92"},
                                            2: {"notes": "not measured"}})
        assert "Temperature remained within" in text
        assert "on all recorded days" not in text
        assert "day(s) 1" in text
