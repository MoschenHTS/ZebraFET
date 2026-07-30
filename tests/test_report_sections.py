"""
test_report_sections.py — Report content that must agree with the analysis.

Covers the sections whose contents were derived independently of the analysis
and could therefore contradict it in the same document: the malformation
appendix (which scored the wrong day and counted dead embryos), the raw-data
appendix (which emitted thousands of rows into one table), the excluded-group
disclosure, and the reference list.
"""
import os

import pandas as pd
import pytest
from docx import Document

from src.core.constants import APP_VERSION
from src.export.report_generator import ReportGenerator

CONC_MAP = {
    "ctrl": {"id": "ctrl", "type": "Control", "value": 0.0, "color": "#4d4d4d"},
    "s1": {"id": "s1", "type": "Substrate", "value": 1.0, "color": "#2166ac"},
}


def _well(status, conditions=()):
    return {"status": status, "sublethal_conditions": list(conditions), "notes": ""}


@pytest.fixture
def snapshot():
    return {
        "project_name": "Sections", "main_researcher": "", "substance": "X",
        "concentration_unit": "mg/L", "start_date": "2025-01-01",
        "num_days": 2, "num_plates": 1, "plate_format": "96-well",
        "report_notes": "", "substance_details": {}, "test_organisms": {},
        "test_conditions": {}, "methodology": {},
        "concentration_settings": {"concentrations": list(CONC_MAP.values())},
        "plate_layout": {"1": {"A1": "ctrl", "B1": "s1", "B2": "s1"}},
        "concentration_map": CONC_MAP,
        "plate_dimensions": (8, 12),
        "photos_with_metadata": [],
        "well_data": {
            "1": {
                "1": {
                    "A1": _well("Live Embryo"),
                    "B1": _well("Live Embryo", ["Pericardial oedema"]),
                    "B2": _well("Live Embryo"),
                }
            },
            "2": {
                "1": {
                    "A1": _well("Live Embryo"),
                    # Dead but carrying a condition recorded before it died.
                    "B1": _well("Dead Embryo", ["Yolk sac oedema"]),
                    "B2": _well("Live Embryo", ["Spinal curvature (scoliosis)"]),
                }
            },
        },
        "completed_days": [1, 2],
    }


def _analysis(**overrides):
    base = {
        "lc50_results": {"lc50": "Not Calculated", "model_info": {}},
        "noec_loec_results": {},
        "summary_df": pd.DataFrame(),
    }
    base.update(overrides)
    return base


def _render(tmp_path, snapshot, analysis, name="report.docx"):
    out = str(tmp_path / name)
    ok = ReportGenerator(snapshot, str(tmp_path), analysis).generate_report(out)
    assert ok, "report generation failed"
    doc = Document(out)
    text = " ".join(p.text for p in doc.paragraphs)
    cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
    return out, text, cells


class TestMalformationAppendix:
    def test_uses_the_analysed_day_not_the_last_day(self, tmp_path, snapshot):
        """Analysing day 1 must not report day 2's malformations."""
        _, _, cells = _render(tmp_path, snapshot, _analysis(analysis_day=1))
        assert any("Pericardial oedema" in c for c in cells)
        assert not any("Yolk sac oedema" in c for c in cells)

    def test_excludes_dead_embryos(self, tmp_path, snapshot):
        """A dead embryo cannot be scored for malformation, as the analysis holds."""
        _, _, cells = _render(tmp_path, snapshot, _analysis(analysis_day=2))
        assert any("Spinal curvature (scoliosis)" in c for c in cells)
        assert not any("Yolk sac oedema" in c for c in cells)

    def test_reports_the_survivor_count(self, tmp_path, snapshot):
        _, _, cells = _render(tmp_path, snapshot, _analysis(analysis_day=2))
        assert any("N (survivors)" in c for c in cells)


class TestRawDataAppendix:
    def test_writes_a_sibling_csv(self, tmp_path, snapshot):
        out, text, _ = _render(tmp_path, snapshot, _analysis(analysis_day=2))
        csv_path = os.path.splitext(out)[0] + "_RawData.csv"
        assert os.path.isfile(csv_path)
        assert os.path.basename(csv_path) in text

    def test_csv_holds_every_observation(self, tmp_path, snapshot):
        out, _, _ = _render(tmp_path, snapshot, _analysis(analysis_day=2))
        df = pd.read_csv(os.path.splitext(out)[0] + "_RawData.csv")
        assert len(df) == 6  # 3 wells x 2 days
        assert set(df["Plate"]) == {1}

    def test_document_does_not_carry_the_rows(self, tmp_path, snapshot):
        """The observation table is the one that grows without bound."""
        out, _, _ = _render(tmp_path, snapshot, _analysis(analysis_day=2))
        doc = Document(out)
        raw_header = {"Day", "Plate", "Well", "Status", "Malformations", "Notes"}
        for table in doc.tables:
            header = {cell.text for cell in table.rows[0].cells}
            assert not raw_header.issubset(header), (
                "raw observations were emitted into the document instead of the CSV"
            )


class TestUnevaluableGroups:
    def test_excluded_groups_are_named(self, tmp_path, snapshot):
        _, text, _ = _render(
            tmp_path, snapshot, _analysis(analysis_day=2, unevaluable_groups=["s2", "s3"])
        )
        assert "s2" in text and "s3" in text
        assert "excluded from the statistical analysis" in text

    def test_nothing_is_claimed_when_all_groups_are_evaluable(self, tmp_path, snapshot):
        _, text, _ = _render(tmp_path, snapshot, _analysis(analysis_day=2))
        assert "excluded from the statistical analysis" not in text


class TestReferences:
    def test_section_is_present(self, tmp_path, snapshot):
        _, text, _ = _render(tmp_path, snapshot, _analysis(analysis_day=2))
        assert "3. References" in text or "References" in text

    def test_self_citation_carries_the_running_version(self, tmp_path, snapshot):
        _, text, _ = _render(tmp_path, snapshot, _analysis(analysis_day=2))
        assert f"(v{APP_VERSION})" in text

    def test_oecd_guideline_is_always_listed(self, tmp_path, snapshot):
        _, text, _ = _render(tmp_path, snapshot, _analysis(analysis_day=2))
        assert "Test No. 236" in text

    def test_abbott_is_listed_only_when_applied(self, tmp_path, snapshot):
        _, without, _ = _render(tmp_path, snapshot, _analysis(analysis_day=2), "a.docx")
        assert "Abbott, W. S. (1925)" not in without

        _, with_abbott, _ = _render(
            tmp_path, snapshot, _analysis(analysis_day=2, abbott_applied=True), "b.docx"
        )
        assert "Abbott, W. S. (1925)" in with_abbott

    def test_model_selection_reference_only_in_auto_mode(self, tmp_path, snapshot):
        manual = _analysis(
            analysis_day=2,
            lc50_results={"lc50": "1.0", "model_info": {"mode": "manual", "n_free": 4}},
        )
        _, text, _ = _render(tmp_path, snapshot, manual, "c.docx")
        assert "Burnham" not in text

        auto = _analysis(
            analysis_day=2,
            lc50_results={
                "lc50": "1.0",
                "model_info": {"mode": "auto", "n_free": 2, "display_name": "2PL"},
            },
        )
        _, text, _ = _render(tmp_path, snapshot, auto, "d.docx")
        assert "Burnham" in text

    def test_appendices_follow_references(self, tmp_path, snapshot):
        _, text, _ = _render(tmp_path, snapshot, _analysis(analysis_day=2))
        assert "4. Appendices" in text
        assert text.index("References") < text.index("4. Appendices")


class TestValiditySectionAppearsOnlyWhenItHasContent:
    """An absence of data is reported by omitting the section, not by a message.

    A project with no control group and no monitoring log has nothing to say about
    validity. Printing "could not be determined" under a heading states a finding
    about the study that the software is in no position to make; an empty heading
    reads as a section that failed to render.
    """

    CONTROLLED = pd.DataFrame([
        {"conc_id": "ctrl", "conc_type": "Control", "conc_value": 0.0,
         "total": 20, "n_scored": 20, "dead": 1, "hatched": 18, "live": 19,
         "malformed": 0},
    ])

    def test_section_is_absent_when_nothing_can_be_assessed(self, tmp_path, snapshot):
        _, text, _ = _render(tmp_path, snapshot, _analysis(analysis_day=2))
        assert "Test Validity Criteria" not in text

    def test_lethal_effects_takes_the_first_subsection_number(self, tmp_path, snapshot):
        """Headings must stay contiguous when validity drops out."""
        out, _, _ = _render(tmp_path, snapshot, _analysis(analysis_day=2))
        headings = [p.text for p in Document(out).paragraphs
                    if p.style.name.startswith("Heading")]
        assert any(h.startswith("2.1 Lethal Effects") for h in headings)
        assert not any("2.2 Lethal Effects" in h for h in headings)

    def test_section_returns_once_a_control_group_exists(self, tmp_path, snapshot):
        out, text, _ = _render(
            tmp_path, snapshot, _analysis(analysis_day=2, summary_df=self.CONTROLLED)
        )
        assert "Test Validity Criteria" in text
        headings = [p.text for p in Document(out).paragraphs
                    if p.style.name.startswith("Heading")]
        assert any(h.startswith("2.1 Test Validity") for h in headings)
        assert any(h.startswith("2.2 Lethal Effects") for h in headings)


class TestReferenceCompleteness:
    """Every method the report names must resolve to a listed source.

    A citation with no entry is a dangling reference; a method named with no
    citation at all is worse, since the reader has nothing to follow.
    """

    CONTROLLED = pd.DataFrame([
        {"conc_id": "ctrl", "conc_type": "Control", "conc_value": 0.0,
         "total": 20, "n_scored": 20, "dead": 1, "hatched": 18, "live": 19,
         "malformed": 0},
        {"conc_id": "s1", "conc_type": "Substrate", "conc_value": 1.0,
         "total": 20, "n_scored": 20, "dead": 8, "hatched": 12, "live": 12,
         "malformed": 3},
    ])

    def _text(self, tmp_path, snapshot, **overrides):
        analysis = _analysis(analysis_day=2, summary_df=self.CONTROLLED, **overrides)
        _, text, _ = _render(tmp_path, snapshot, analysis)
        return text

    def test_fishers_exact_test_is_cited(self, tmp_path, snapshot):
        text = self._text(tmp_path, snapshot)
        assert "Fisher's Exact Test (Fisher, 1922)" in text
        assert "Fisher, R. A. (1922)" in text

    def test_the_guideline_citation_is_the_current_edition(self, tmp_path, snapshot):
        text = self._text(tmp_path, snapshot)
        assert "OECD (2025)" in text
        assert "OECD (2013)" not in text

    def test_the_haldane_anscombe_correction_is_cited(self, tmp_path, snapshot):
        text = self._text(tmp_path, snapshot)
        assert "Haldane, 1956; Anscombe, 1956" in text
        assert "Haldane, J. B. S. (1956)" in text
        assert "Anscombe, F. J. (1956)" in text

    def test_the_bootstrap_and_the_model_are_cited(self, tmp_path, snapshot):
        text = self._text(tmp_path, snapshot, lc50_results={
            "lc50": "1.0", "model_info": {"display_name": "4PL (all free)", "n_free": 4},
            "_fitted_params": [0.0, 100.0, -2.0, 1.0],
        })
        assert "Efron, 1979" in text and "Efron, B. (1979)" in text
        assert "Ritz, 2010" in text and "Ritz, C. (2010)" in text

    def test_bonferroni_is_cited_when_it_is_the_correction_used(self, tmp_path, snapshot):
        """Only Holm was ever credited, so a Bonferroni run cited nothing."""
        text = self._text(tmp_path, snapshot, noec_loec_results={
            "noec": "1.0000", "loec": "Not detected", "correction_label": "Bonferroni",
        })
        assert "Bonferroni" in text
        assert "Dunn, O. J. (1961)" in text

    def test_holm_is_cited_instead_when_holm_was_used(self, tmp_path, snapshot):
        text = self._text(tmp_path, snapshot, noec_loec_results={
            "noec": "1.0000", "loec": "Not detected", "correction_label": "Holm-Bonferroni",
        })
        assert "Holm, S. (1979)" in text
        assert "Dunn, O. J. (1961)" not in text

    def test_the_validity_claim_names_more_than_control_mortality(self, tmp_path, snapshot):
        """The Methods section described one criterion of six as the requirement."""
        text = self._text(tmp_path, snapshot)
        assert "contingent upon the mortality in the control group(s) not exceeding" not in text
        assert "control hatching rate at the end of the exposure" in text
