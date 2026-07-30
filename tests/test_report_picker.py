"""
test_report_picker.py — The report can be generated in part.

The full document runs to every appendix and figure, which suits a final study
record but not the interim reads that happen far more often. Omitting a section
must renumber what remains: a document whose headings jump from 1 to 3, or whose
appendices run A, C, D, reads as though pages went missing.
"""
import re

import pandas as pd
import pytest
from docx import Document

from src.export.report_generator import ReportGenerator

CONC_MAP = {
    "ctrl": {"id": "ctrl", "type": "Control", "value": 0.0, "color": "#4d4d4d"},
    "s1": {"id": "s1", "type": "Substrate", "value": 1.0, "color": "#2166ac"},
}


@pytest.fixture
def snapshot():
    return {
        "project_name": "Picker", "main_researcher": "", "substance": "X",
        "concentration_unit": "mg/L", "start_date": "2025-01-01",
        "num_days": 1, "num_plates": 1, "plate_format": "96-well",
        "report_notes": "", "substance_details": {}, "test_organisms": {},
        "test_conditions": {}, "methodology": {},
        "concentration_settings": {"concentrations": list(CONC_MAP.values())},
        "plate_layout": {"1": {"A1": "ctrl", "B1": "s1"}},
        "concentration_map": CONC_MAP,
        "plate_dimensions": (8, 12),
        "photos_with_metadata": [],
        "well_data": {
            "1": {"1": {
                "A1": {"status": "Live Embryo", "sublethal_conditions": [], "notes": ""},
                "B1": {"status": "Live Embryo", "sublethal_conditions": ["Yolk sac oedema"], "notes": ""},
            }}
        },
        "completed_days": [1],
    }


@pytest.fixture
def analysis():
    return {
        "analysis_day": 1,
        "lc50_results": {"lc50": "Not Calculated", "model_info": {}},
        "noec_loec_results": {},
        # A control group with scored embryos, so the validity section has
        # something to report; it is omitted entirely when nothing can be
        # assessed, which would make the picker tests below vacuous.
        "summary_df": pd.DataFrame([
            {"conc_id": "ctrl", "conc_type": "Control", "conc_value": 0.0,
             "total": 20, "n_scored": 20, "dead": 1, "hatched": 18, "live": 19,
             "malformed": 0},
            {"conc_id": "s1", "conc_type": "Substrate", "conc_value": 1.0,
             "total": 20, "n_scored": 20, "dead": 5, "hatched": 14, "live": 15,
             "malformed": 2},
        ]),
        "abbott_applied": True,  # pulls Abbott into the reference list
    }


def _render(tmp_path, snapshot, analysis, sections=None, name="r.docx"):
    out = str(tmp_path / name)
    ok = ReportGenerator(snapshot, str(tmp_path), analysis, sections=sections).generate_report(out)
    assert ok, "report generation failed"
    doc = Document(out)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    text = " ".join(p.text for p in doc.paragraphs)
    return headings, text


def _top_level(headings):
    return [h for h in headings if re.match(r"^\d+\. ", h)]


class TestDefaultIsEverything:
    def test_omitting_sections_argument_keeps_all(self, tmp_path, snapshot, analysis):
        headings, _ = _render(tmp_path, snapshot, analysis)
        titles = [h.split(". ", 1)[1] for h in _top_level(headings)]
        assert titles == ["Materials and Methods", "Results", "References", "Appendices"]

    def test_numbering_is_contiguous_from_one(self, tmp_path, snapshot, analysis):
        headings, _ = _render(tmp_path, snapshot, analysis)
        numbers = [int(h.split(".", 1)[0]) for h in _top_level(headings)]
        assert numbers == [1, 2, 3, 4]


class TestOmittingSections:
    def test_dropping_methods_renumbers_the_rest(self, tmp_path, snapshot, analysis):
        sections = ReportGenerator.ALL_SECTIONS - {"methods"}
        headings, _ = _render(tmp_path, snapshot, analysis, sections)
        top = _top_level(headings)
        assert top[0].startswith("1. Results")
        assert [int(h.split(".", 1)[0]) for h in top] == [1, 2, 3]

    def test_results_subsections_follow_the_section_number(self, tmp_path, snapshot, analysis):
        sections = ReportGenerator.ALL_SECTIONS - {"methods"}
        headings, _ = _render(tmp_path, snapshot, analysis, sections)
        assert any(h.startswith("1.1 Test Validity") for h in headings)
        assert not any(h.startswith("2.1 Test Validity") for h in headings)

    def test_references_cannot_be_dropped(self, tmp_path, snapshot, analysis):
        """The methods text cites inline, and ZebraFET cites itself from that list.

        Offering References as optional left in-text citations pointing at nothing
        and removed the software's own citation along with the section.
        """
        assert "references" not in ReportGenerator.ALL_SECTIONS
        headings, text = _render(tmp_path, snapshot, analysis, sections=set())
        assert any("References" in h for h in headings)
        assert "ZebraFET" in text

    def test_dropping_every_appendix_removes_the_heading(self, tmp_path, snapshot, analysis):
        sections = {s for s in ReportGenerator.ALL_SECTIONS if not s.startswith("appendix_")}
        headings, _ = _render(tmp_path, snapshot, analysis, sections)
        assert not any("Appendices" in h for h in headings)

    def test_core_results_survive_every_omission(self, tmp_path, snapshot, analysis):
        """Validity and lethal endpoints are not optional."""
        headings, _ = _render(tmp_path, snapshot, analysis, sections=set())
        assert any("Test Validity Criteria" in h for h in headings)
        assert any("Lethal Effects" in h for h in headings)


class TestAppendixLettering:
    def test_letters_are_contiguous_when_one_is_dropped(self, tmp_path, snapshot, analysis):
        """A, C, D would read as a missing appendix rather than an omitted one."""
        sections = ReportGenerator.ALL_SECTIONS - {"appendix_raw"}
        headings, _ = _render(tmp_path, snapshot, analysis, sections)
        letters = [h.split(":")[0] for h in headings if h.startswith("Appendix ")]
        assert letters == ["Appendix A", "Appendix B", "Appendix C"]

    def test_remaining_appendices_keep_their_titles(self, tmp_path, snapshot, analysis):
        sections = ReportGenerator.ALL_SECTIONS - {"appendix_layout"}
        headings, _ = _render(tmp_path, snapshot, analysis, sections)
        appendices = [h for h in headings if h.startswith("Appendix ")]
        assert appendices[0] == "Appendix A: Raw Observational Data"


class TestReferencesFollowContent:
    def test_a_citation_cannot_outlive_its_section(self, tmp_path, snapshot, analysis):
        """Abbott is cited by the methods text, so dropping methods drops the work."""
        _, with_methods = _render(tmp_path, snapshot, analysis, name="a.docx")
        assert "Abbott, W. S. (1925)" in with_methods

        sections = ReportGenerator.ALL_SECTIONS - {"methods"}
        _, without = _render(tmp_path, snapshot, analysis, sections, name="b.docx")
        assert "Abbott, W. S. (1925)" not in without

    def test_the_guideline_is_still_listed_without_methods(self, tmp_path, snapshot, analysis):
        sections = ReportGenerator.ALL_SECTIONS - {"methods"}
        _, text = _render(tmp_path, snapshot, analysis, sections)
        assert "ZebraFET" in text
