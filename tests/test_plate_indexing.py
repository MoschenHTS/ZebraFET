"""
test_plate_indexing.py — Plate indices are 1-based end to end.

The UI writes plate_index starting at 1 (see the plate loops in
PlateLayoutPage and ExperimentViewWidget), so the stored value is already the
number an operator sees on the plate tab. Display code that adds 1 to reach a
human-readable number reports every plate as its successor: a six-plate study
rendered as "Plate 2" through "Plate 7" in every generated report.

These tests pin the convention at both ends — what ProjectManager stores, and
what the report renders — so the two cannot drift apart again.
"""
import pandas as pd
import pytest
from docx import Document

from src.core.constants import FIRST_PLATE_INDEX
from src.core.project_manager import ProjectManager
from src.export.report_generator import ReportGenerator

CONCENTRATIONS = [
    {"id": "ctrl", "type": "Control", "value": 0.0, "replicates": 1, "wells": 2, "per_plate": True},
    {"id": "s1", "type": "Substrate", "value": 1.0, "replicates": 1, "wells": 2, "per_plate": True},
]


class TestStorageConvention:
    def test_first_plate_index_is_one(self):
        assert FIRST_PLATE_INDEX == 1

    def test_committed_layout_keeps_its_plate_numbers(self, tmp_path):
        """commit_plate_layout must store the key it was given, not an offset."""
        m = ProjectManager.create_new(
            str(tmp_path / "P"), {"project_name": "P", "num_days": 1, "num_plates": 2}
        )
        try:
            m.set_concentrations(CONCENTRATIONS, required_embryos=4, required_plates=2)
            m.commit_plate_layout({"1": {"A1": "ctrl"}, "2": {"A1": "s1"}})
            stored = [
                r[0] for r in m._conn.execute(
                    "SELECT DISTINCT plate_index FROM plate_layout ORDER BY plate_index"
                ).fetchall()
            ]
            assert stored == [1, 2]
            assert m.get_plate_layout(1) == {"A1": "ctrl"}
            assert m.get_plate_layout(2) == {"A1": "s1"}
        finally:
            m.close()

    def test_observations_round_trip_under_their_own_plate_number(self, tmp_path):
        m = ProjectManager.create_new(
            str(tmp_path / "P"), {"project_name": "P", "num_days": 1, "num_plates": 1}
        )
        try:
            m.set_concentrations(CONCENTRATIONS, required_embryos=4, required_plates=1)
            m.commit_plate_layout({"1": {"A1": "ctrl"}})
            m.save_well_data(1, 1, "A1", "Dead Embryo", [], [], "")
            assert m.get_well_data(1, 1, "A1")["status"] == "Dead Embryo"
            assert "1" in m.get_well_observations_for_day(1)
        finally:
            m.close()


class TestReportRendersStoredPlateNumbers:
    @pytest.fixture
    def two_plate_snapshot(self):
        conc_map = {
            "ctrl": {"id": "ctrl", "type": "Control", "value": 0.0, "color": "#4d4d4d"},
            "s1": {"id": "s1", "type": "Substrate", "value": 1.0, "color": "#2166ac"},
        }
        return {
            "project_name": "Indexing", "main_researcher": "", "substance": "X",
            "concentration_unit": "mg/L", "start_date": "2025-01-01",
            "num_days": 1, "num_plates": 2, "plate_format": "96-well",
            "report_notes": "", "substance_details": {}, "test_organisms": {},
            "test_conditions": {}, "methodology": {},
            "concentration_settings": {"concentrations": list(conc_map.values())},
            "plate_layout": {"1": {"A1": "ctrl"}, "2": {"A1": "s1"}},
            "concentration_map": conc_map,
            "plate_dimensions": (8, 12),
            "photos_with_metadata": [],
            "well_data": {
                "1": {
                    "1": {"A1": {"status": "Live Embryo", "sublethal_conditions": [], "notes": ""}},
                    "2": {"A1": {"status": "Live Embryo", "sublethal_conditions": [], "notes": ""}},
                }
            },
            "completed_days": [1],
        }

    @pytest.fixture
    def analysis(self):
        return {
            "lc50_results": {"lc50": "Not Calculated", "model_info": {}},
            "noec_loec_results": {},
            "summary_df": pd.DataFrame(),
            "analysis_day": 1,
        }

    def _document_text(self, tmp_path, snapshot, analysis):
        out = str(tmp_path / "report.docx")
        ReportGenerator(snapshot, str(tmp_path), analysis).generate_report(out)
        doc = Document(out)
        paragraphs = [p.text for p in doc.paragraphs]
        cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
        return " ".join(paragraphs + cells)

    def test_plate_captions_match_the_stored_numbers(self, tmp_path, two_plate_snapshot, analysis):
        text = self._document_text(tmp_path, two_plate_snapshot, analysis)
        assert "Plate 1" in text
        assert "Plate 2" in text

    def test_no_plate_is_reported_as_its_successor(self, tmp_path, two_plate_snapshot, analysis):
        """The off-by-one showed up as a plate number one past the last plate."""
        text = self._document_text(tmp_path, two_plate_snapshot, analysis)
        assert "Plate 3" not in text
        assert "Plate 0" not in text
