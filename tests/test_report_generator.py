"""
test_report_generator.py — Structural and functional tests for ReportGenerator.

Verifies that generate_report() produces a valid .docx with the expected document
structure (headings, tables, validity text) using a snapshot dict — no database
or Qt dependency required.
"""
import os

import pandas as pd
import pytest
from docx import Document

from src.export.report_generator import ReportGenerator


@pytest.fixture
def minimal_snapshot():
    return {
        "project_name": "Test Project",
        "main_researcher": "Dr. Test",
        "substance": "TestChem",
        "concentration_unit": "mg/L",
        "start_date": "2025-01-01",
        "num_days": 4,
        "num_plates": 1,
        "plate_format": "96-well",
        "report_notes": "",
        "substance_details": {
            "cas_number": "1234-56-7",
            "molecular_weight": "200",
            "purity": "99",
            "supplier": "Sigma",
            "physical_appearance": "White powder",
            "water_solubility": "500 mg/L",
            "iupac_name": "test-iupac",
            "solvent_used": "",
            "positive_control_substance": "",
        },
        "test_organisms": {
            "strain": "AB",
            "source": "In-house",
            "collection_method": "Standard protocol",
        },
        "test_conditions": {
            "water_type": "ISO water",
            "temperature": "26",
            "ph": "7.2",
            "hardness": "250",
            "conductivity": "500",
            "dissolved_oxygen": "8",
            "photoperiod": "16:8",
            "acceptable_mortality": 10.0,
        },
        "methodology": {
            "test_procedure": "Static",
            "solution_preparation": "Serial dilution",
            "selection_criteria": "Standard",
        },
        "concentration_settings": {
            "concentrations": [
                {"id": "ctrl", "type": "Control",   "value": 0.0, "replicates": 1, "wells": 4, "per_plate": True, "color": "#4d4d4d", "sort_order": 0},
                {"id": "s1",   "type": "Substrate", "value": 1.0, "replicates": 1, "wells": 4, "per_plate": True, "color": "#2166ac", "sort_order": 1},
                {"id": "s2",   "type": "Substrate", "value": 2.0, "replicates": 1, "wells": 4, "per_plate": True, "color": "#2166ac", "sort_order": 2},
            ],
            "required_embryos": 20,
            "required_plates": 1,
        },
        "plate_layout": {
            "1": {
                "A1": "ctrl", "A2": "ctrl",
                "B1": "s1",   "B2": "s1",
                "C1": "s2",   "C2": "s2",
            }
        },
        "concentration_map": {
            "ctrl": {"id": "ctrl", "type": "Control",   "value": 0.0, "replicates": 1, "wells": 4, "per_plate": True, "color": "#4d4d4d", "sort_order": 0},
            "s1":   {"id": "s1",   "type": "Substrate", "value": 1.0, "replicates": 1, "wells": 4, "per_plate": True, "color": "#2166ac", "sort_order": 1},
            "s2":   {"id": "s2",   "type": "Substrate", "value": 2.0, "replicates": 1, "wells": 4, "per_plate": True, "color": "#2166ac", "sort_order": 2},
        },
        "plate_dimensions": (8, 12),
        "photos_with_metadata": [],
        "well_data": {},
        "completed_days": [],
    }


@pytest.fixture
def minimal_analysis():
    return {
        "lc50_results": {
            "lc50": "Not Calculated",
            "slope": "Not Calculated",
            "r_squared": "Not Calculated",
            "model_info": {
                "display_name": "4PL (all free)",
                "mode": "manual",
                "bottom": None,
                "top": None,
                "n_free": 4,
                "aic_table": None,
            },
        },
        "noec_loec_results": {"noec": "Not Calculated", "loec": "Not Calculated"},
        "summary_df": pd.DataFrame(),
        "mortality_plot_figure": None,
        "hatching_plot_figure": None,
        "malformation_plot_figure": None,
    }


@pytest.fixture
def full_analysis():
    """Analysis results with a calculated LC50 and non-empty summary table."""
    summary_df = pd.DataFrame([
        {"conc_id": "ctrl", "conc_type": "Control",   "conc_value": 0.0, "total": 4, "dead": 0, "hatched": 2, "malformed": 0},
        {"conc_id": "s1",   "conc_type": "Substrate", "conc_value": 1.0, "total": 4, "dead": 1, "hatched": 1, "malformed": 0},
        {"conc_id": "s2",   "conc_type": "Substrate", "conc_value": 2.0, "total": 4, "dead": 3, "hatched": 0, "malformed": 1},
    ])
    return {
        "lc50_results": {
            "lc50": "1.5000 (95% CI: 1.2000 – 1.8000)",
            "slope": "-2.5000",
            "r_squared": "0.9800",
            "model_info": {
                "display_name": "2PL (bottom=0.0%, top=100.0%)",
                "mode": "manual",
                "bottom": 0.0,
                "top": 100.0,
                "n_free": 2,
                "aic_table": None,
            },
            "_fitted_params": [0.0, 100.0, -2.5, 1.5],
        },
        "noec_loec_results": {"noec": "1.0000", "loec": "2.0000"},
        "summary_df": summary_df,
        "mortality_plot_figure": None,
        "hatching_plot_figure": None,
        "malformation_plot_figure": None,
    }


class TestReportGeneratorStructure:
    def test_generate_report_returns_true(self, tmp_path, minimal_snapshot, minimal_analysis):
        out = str(tmp_path / "report.docx")
        assert ReportGenerator(minimal_snapshot, str(tmp_path), minimal_analysis).generate_report(out)

    def test_output_file_is_created(self, tmp_path, minimal_snapshot, minimal_analysis):
        out = str(tmp_path / "report.docx")
        ReportGenerator(minimal_snapshot, str(tmp_path), minimal_analysis).generate_report(out)
        assert os.path.isfile(out)

    def test_document_has_materials_and_methods_heading(self, tmp_path, minimal_snapshot, minimal_analysis):
        out = str(tmp_path / "report.docx")
        ReportGenerator(minimal_snapshot, str(tmp_path), minimal_analysis).generate_report(out)
        headings = [p.text for p in Document(out).paragraphs if "Heading" in p.style.name]
        assert any("Materials and Methods" in h for h in headings)

    def test_document_has_results_heading(self, tmp_path, minimal_snapshot, minimal_analysis):
        out = str(tmp_path / "report.docx")
        ReportGenerator(minimal_snapshot, str(tmp_path), minimal_analysis).generate_report(out)
        headings = [p.text for p in Document(out).paragraphs if "Heading" in p.style.name]
        assert any("Results" in h for h in headings)

    def test_document_has_appendices_heading(self, tmp_path, minimal_snapshot, minimal_analysis):
        out = str(tmp_path / "report.docx")
        ReportGenerator(minimal_snapshot, str(tmp_path), minimal_analysis).generate_report(out)
        headings = [p.text for p in Document(out).paragraphs if "Heading" in p.style.name]
        assert any("Appendix" in h or "Appendices" in h for h in headings)

    def test_document_has_substance_name_in_title(self, tmp_path, minimal_snapshot, minimal_analysis):
        out = str(tmp_path / "report.docx")
        ReportGenerator(minimal_snapshot, str(tmp_path), minimal_analysis).generate_report(out)
        all_text = " ".join(p.text for p in Document(out).paragraphs)
        assert "TestChem" in all_text

    def test_minimum_table_count(self, tmp_path, minimal_snapshot, minimal_analysis):
        out = str(tmp_path / "report.docx")
        ReportGenerator(minimal_snapshot, str(tmp_path), minimal_analysis).generate_report(out)
        # M&M: substance (1), organisms (2), conditions (3) + plate layout appendix (4)
        assert len(Document(out).tables) >= 4

    def test_no_photos_produces_correct_message(self, tmp_path, minimal_snapshot, minimal_analysis):
        out = str(tmp_path / "report.docx")
        ReportGenerator(minimal_snapshot, str(tmp_path), minimal_analysis).generate_report(out)
        all_text = " ".join(p.text for p in Document(out).paragraphs)
        assert "No photographic documentation" in all_text

    def test_full_analysis_adds_results_and_curve_tables(self, tmp_path, minimal_snapshot, full_analysis):
        out = str(tmp_path / "report.docx")
        ReportGenerator(minimal_snapshot, str(tmp_path), full_analysis).generate_report(out)
        assert len(Document(out).tables) >= 5

    def test_lc50_value_present_in_document_when_calculated(self, tmp_path, minimal_snapshot, full_analysis):
        out = str(tmp_path / "report.docx")
        ReportGenerator(minimal_snapshot, str(tmp_path), full_analysis).generate_report(out)
        all_text = " ".join(p.text for p in Document(out).paragraphs)
        assert "1.5000" in all_text


class TestReportFormatting:
    """Publication-facing formatting the coworker/reviewer actually reads."""

    def _summary_headers(self, tmp_path, snapshot, analysis):
        out = str(tmp_path / "report.docx")
        ReportGenerator(snapshot, str(tmp_path), analysis).generate_report(out)
        for table in Document(out).tables:
            headers = [c.text for c in table.rows[0].cells]
            if "Mortality (%)" in headers:
                return headers
        return []

    def test_summary_table_dead_header_is_capitalized(self, tmp_path, minimal_snapshot, full_analysis):
        """The 'dead' count column must render title-cased like its siblings."""
        headers = self._summary_headers(tmp_path, minimal_snapshot, full_analysis)
        assert "Dead" in headers
        assert "dead" not in headers

    def test_single_replicate_is_singular(self, tmp_path, minimal_snapshot, minimal_analysis):
        out = str(tmp_path / "report.docx")
        ReportGenerator(minimal_snapshot, str(tmp_path), minimal_analysis).generate_report(out)
        all_text = " ".join(p.text for p in Document(out).paragraphs)
        assert "1 replicate for each treatment group" in all_text
        assert "1 replicates" not in all_text

    def test_multiple_replicates_stay_plural(self, tmp_path, minimal_snapshot, minimal_analysis):
        snap = {**minimal_snapshot}
        concs = [dict(c, replicates=3) for c in snap["concentration_settings"]["concentrations"]]
        snap["concentration_settings"] = {**snap["concentration_settings"], "concentrations": concs}
        out = str(tmp_path / "report.docx")
        ReportGenerator(snap, str(tmp_path), minimal_analysis).generate_report(out)
        all_text = " ".join(p.text for p in Document(out).paragraphs)
        assert "3 replicates for each treatment group" in all_text


class TestPhotoPanelLabelling:
    """The photo panel letters must stay in lock-step with the returned metadata
    even when an image fails to load (otherwise the appendix legend desyncs)."""

    def _make_image(self, path):
        from PIL import Image
        os.makedirs(os.path.dirname(path), exist_ok=True)
        Image.new("RGB", (60, 40), (120, 160, 200)).save(path)

    def test_unloadable_image_is_dropped_from_returned_meta(self, tmp_path, minimal_snapshot, minimal_analysis):
        proj_dir = str(tmp_path)
        self._make_image(os.path.join(proj_dir, "photos", "Day_1", "good1.png"))
        self._make_image(os.path.join(proj_dir, "photos", "Day_1", "good2.png"))
        photos_meta = [
            {"path": "photos/Day_1/good1.png", "day": 1, "plate": 1, "well": "A1"},
            {"path": "photos/Day_1/missing.png", "day": 1, "plate": 1, "well": "B1"},
            {"path": "photos/Day_1/good2.png", "day": 1, "plate": 1, "well": "C1"},
        ]
        gen = ReportGenerator(minimal_snapshot, proj_dir, minimal_analysis)
        buffer, returned = gen._create_photo_panel(photos_meta)
        assert buffer is not None
        # Only the two loadable images survive, in order, with no gap.
        assert [m["well"] for m in returned] == ["A1", "C1"]

    def test_all_images_valid_returns_all_meta(self, tmp_path, minimal_snapshot, minimal_analysis):
        proj_dir = str(tmp_path)
        self._make_image(os.path.join(proj_dir, "photos", "Day_1", "a.png"))
        self._make_image(os.path.join(proj_dir, "photos", "Day_1", "b.png"))
        photos_meta = [
            {"path": "photos/Day_1/a.png", "day": 1, "plate": 1, "well": "A1"},
            {"path": "photos/Day_1/b.png", "day": 1, "plate": 1, "well": "B1"},
        ]
        gen = ReportGenerator(minimal_snapshot, proj_dir, minimal_analysis)
        buffer, returned = gen._create_photo_panel(photos_meta)
        assert buffer is not None
        assert [m["well"] for m in returned] == ["A1", "B1"]

    def test_noec_loec_present_in_document_when_calculated(self, tmp_path, minimal_snapshot, full_analysis):
        out = str(tmp_path / "report.docx")
        ReportGenerator(minimal_snapshot, str(tmp_path), full_analysis).generate_report(out)
        all_text = " ".join(p.text for p in Document(out).paragraphs)
        assert "NOEC" in all_text
        assert "LOEC" in all_text

    def test_plate_layout_table_uses_snapshot_data(self, tmp_path, minimal_snapshot, minimal_analysis):
        out = str(tmp_path / "report.docx")
        ReportGenerator(minimal_snapshot, str(tmp_path), minimal_analysis).generate_report(out)
        doc = Document(out)
        all_cells = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
        assert any("ctrl" in c for c in all_cells)


class TestPhotoAppendix:
    """Tests for the TIFF conversion and photo panel compositor."""

    @pytest.fixture
    def snapshot_with_tiff(self, tmp_path, minimal_snapshot):
        """Snapshot wired to a real TIFF placed inside the project directory tree."""
        from PIL import Image

        photo_dir = tmp_path / "photos" / "Day_1"
        photo_dir.mkdir(parents=True)
        tif_path = photo_dir / "A1_Plate1_test.tif"

        # Synthetic RGB TIFF — same mode and format as a real microscopy image
        img = Image.new("RGB", (320, 240), color=(40, 80, 120))
        img.save(str(tif_path), format="TIFF")

        snap = dict(minimal_snapshot)
        snap["photos_with_metadata"] = [
            {"path": "photos/Day_1/A1_Plate1_test.tif", "day": 1, "plate": 1, "well": "A1"}
        ]
        snap["well_data"] = {
            "1": {"1": {"A1": {"status": "Live Embryo", "sublethal_conditions": [], "notes": ""}}}
        }
        return snap

    def test_generate_report_succeeds_with_tiff_photo(self, tmp_path, snapshot_with_tiff, minimal_analysis):
        out = str(tmp_path / "report.docx")
        assert ReportGenerator(snapshot_with_tiff, str(tmp_path), minimal_analysis).generate_report(out)

    def test_photo_appendix_replaces_no_photo_message(self, tmp_path, snapshot_with_tiff, minimal_analysis):
        out = str(tmp_path / "report.docx")
        ReportGenerator(snapshot_with_tiff, str(tmp_path), minimal_analysis).generate_report(out)
        all_text = " ".join(p.text for p in Document(out).paragraphs)
        assert "No photographic documentation" not in all_text

    def test_photo_appendix_includes_day_caption(self, tmp_path, snapshot_with_tiff, minimal_analysis):
        out = str(tmp_path / "report.docx")
        ReportGenerator(snapshot_with_tiff, str(tmp_path), minimal_analysis).generate_report(out)
        all_text = " ".join(p.text for p in Document(out).paragraphs)
        assert "Day 1" in all_text

    def test_tiff_conversion_does_not_leave_temp_files(self, tmp_path, snapshot_with_tiff, minimal_analysis):
        import tempfile
        tmp_before = set(os.listdir(tempfile.gettempdir()))
        out = str(tmp_path / "report.docx")
        ReportGenerator(snapshot_with_tiff, str(tmp_path), minimal_analysis).generate_report(out)
        tmp_after = set(os.listdir(tempfile.gettempdir()))
        new_tmp_pngs = [f for f in (tmp_after - tmp_before) if f.endswith(".png")]
        assert not new_tmp_pngs


class TestGetTestValidityMessage:
    def _make_df(self, control_dead, control_total, treatments=None):
        rows = [{"conc_id": "ctrl", "conc_type": "Control", "conc_value": 0.0,
                 "total": control_total, "dead": control_dead}]
        for i, (cv, d, t) in enumerate(treatments or []):
            rows.append({"conc_id": f"s{i}", "conc_type": "Substrate",
                         "conc_value": cv, "total": t, "dead": d})
        return pd.DataFrame(rows)

    def test_valid_test_at_zero_mortality(self):
        df = self._make_df(0, 10)
        msg = ReportGenerator.get_test_validity_message(df, {"test_conditions": {"acceptable_mortality": 10.0}})
        assert "valid" in msg.lower()
        assert "INVALID" not in msg

    def test_valid_test_at_exact_threshold(self):
        df = self._make_df(1, 10)  # 10.0%
        msg = ReportGenerator.get_test_validity_message(df, {"test_conditions": {"acceptable_mortality": 10.0}})
        assert "valid" in msg.lower()
        assert "INVALID" not in msg

    def test_invalid_test_above_threshold(self):
        df = self._make_df(2, 10)  # 20%
        msg = ReportGenerator.get_test_validity_message(df, {"test_conditions": {"acceptable_mortality": 10.0}})
        assert "INVALID" in msg

    def test_custom_threshold_respected(self):
        df = self._make_df(1, 10)  # 10% > 5% threshold
        msg = ReportGenerator.get_test_validity_message(df, {"test_conditions": {"acceptable_mortality": 5.0}})
        assert "INVALID" in msg

    def test_empty_dataframe_yields_no_statement(self):
        """Nothing to assess is reported by omission, not by a message saying so."""
        assert ReportGenerator.get_test_validity_message(pd.DataFrame(), {}) is None

    def test_no_control_group_yields_no_statement(self):
        df = pd.DataFrame([{"conc_id": "s1", "conc_type": "Substrate",
                             "conc_value": 1.0, "total": 10, "dead": 2}])
        assert ReportGenerator.get_test_validity_message(df, {}) is None

    def test_solvent_control_counts_as_control(self):
        rows = [
            {"conc_id": "sc", "conc_type": "Solvent Control", "conc_value": 0.0, "total": 10, "dead": 0},
            {"conc_id": "s1", "conc_type": "Substrate",       "conc_value": 1.0, "total": 10, "dead": 2},
        ]
        msg = ReportGenerator.get_test_validity_message(
            pd.DataFrame(rows), {"test_conditions": {"acceptable_mortality": 10.0}}
        )
        assert "valid" in msg.lower()
        assert "INVALID" not in msg

    def test_mortality_percentage_included_in_message(self):
        df = self._make_df(0, 10)
        msg = ReportGenerator.get_test_validity_message(df, {"test_conditions": {"acceptable_mortality": 10.0}})
        assert "0.00%" in msg

    def test_control_and_solvent_control_combined(self):
        rows = [
            {"conc_id": "ctrl", "conc_type": "Control",        "conc_value": 0.0, "total": 10, "dead": 1},
            {"conc_id": "sc",   "conc_type": "Solvent Control", "conc_value": 0.0, "total": 10, "dead": 1},
        ]
        msg = ReportGenerator.get_test_validity_message(
            pd.DataFrame(rows), {"test_conditions": {"acceptable_mortality": 10.0}}
        )
        # Combined: 2 dead / 20 total = 10% → still valid
        assert "valid" in msg.lower()
        assert "INVALID" not in msg


class TestValidityVerdictClaimsOnlyWhatItChecked:
    """The verdict is scoped to the criteria ZebraFET can actually evaluate.

    Control mortality alone once produced an unqualified "the test is considered
    valid … specified by OECD TG 236" — a full-guideline claim from one criterion,
    which passed a test whose dissolved oxygen never reached the guideline minimum.
    """

    def _controls(self, dead=1, total=20, hatched=18):
        return pd.DataFrame([{
            "conc_id": "ctrl", "conc_type": "Control", "conc_value": 0.0,
            "total": total, "n_scored": total, "dead": dead, "hatched": hatched,
            "live": total - dead, "malformed": 0,
        }])

    def _msg(self, df=None, **kwargs):
        return ReportGenerator.get_test_validity_message(
            self._controls() if df is None else df,
            {"test_conditions": {"acceptable_mortality": 10.0}}, **kwargs
        )

    def test_a_criterion_without_data_is_not_mentioned_at_all(self):
        """An absence is reported by omission, never by announcing it.

        A criterion the software was never given data for is not a finding about
        the study, and naming it reads as something the study failed to do.
        """
        msg = self._msg()
        assert "not evaluated" not in msg.lower()
        assert "could not be determined" not in msg.lower()
        assert "fertilization" not in msg.lower()
        assert "positive control" not in msg.lower()

    def test_hatching_is_ignored_before_the_end_of_the_test(self):
        """Embryos hatch from ~48-72 hpf, so an interim day cannot fail this."""
        msg = self._msg(self._controls(hatched=0), end_of_test=False)
        assert "hatching" not in msg.lower()
        assert "INVALID" not in msg

    def test_low_hatching_at_the_end_of_the_test_invalidates(self):
        msg = self._msg(self._controls(hatched=2), end_of_test=True,
                        timepoint="96 hpf")
        assert "INVALID" in msg
        assert "hatching rate" in msg
        assert "96 hpf" in msg

    def test_adequate_hatching_is_reported_as_met(self):
        msg = self._msg(end_of_test=True)
        assert "INVALID" not in msg
        assert "hatching rate" in msg

    def test_a_frame_without_hatching_counts_is_tolerated(self):
        """Older results dicts carry no 'hatched' column; that must not raise."""
        bare = pd.DataFrame([{"conc_id": "ctrl", "conc_type": "Control",
                              "conc_value": 0.0, "total": 20, "dead": 1}])
        msg = self._msg(bare, end_of_test=True)
        assert "hatching" not in msg.lower()

    def test_out_of_range_water_quality_invalidates_the_test(self):
        """The advisory and the verdict must not disagree about the same reading."""
        findings = [{"key": "dissolved_oxygen", "subject": "Dissolved oxygen",
                     "advisory": "Dissolved-oxygen", "criterion": "the OECD TG 236 minimum "
                     "of 80% saturation", "measured": [3], "out_of_range": ["Day 3 (74%)"],
                     "passed": False, "scope_is_every_day": True}]
        msg = self._msg(water_quality_findings=findings)
        assert "INVALID" in msg
        assert "Day 3 (74%)" in msg

    def test_in_range_water_quality_counts_towards_the_criteria_met(self):
        findings = [{"key": "temperature", "subject": "Temperature",
                     "advisory": "Temperature", "criterion": "the OECD TG 236 range of "
                     "26 ± 1 °C", "measured": [1], "out_of_range": [],
                     "passed": True, "scope_is_every_day": True}]
        msg = self._msg(water_quality_findings=findings)
        assert "INVALID" not in msg
        assert "all 2" in msg


class TestReportStatisticsSections:
    """v2.2.0 report additions: sublethal battery, teratogenic index, effect sizes."""

    def _analysis(self):
        from collections import Counter
        from src.core.biostatistics import (sublethal_endpoint_stats,
                                            calculate_teratogenic_index, calculate_lc50_robust,
                                            pooled_control_mortality_pct, abbott_correct)
        summary_df = pd.DataFrame([
            {"conc_id": "ctrl", "conc_type": "Control",  "conc_value": 0.0, "total": 20, "dead": 1,
             "hatched": 16, "malformed": 0, "malformation_details": Counter()},
            {"conc_id": "s1", "conc_type": "Substrate", "conc_value": 0.5, "total": 20, "dead": 2,
             "hatched": 15, "malformed": 2, "malformation_details": Counter({"Pericardial oedema": 2})},
            {"conc_id": "s2", "conc_type": "Substrate", "conc_value": 1.0, "total": 20, "dead": 4,
             "hatched": 13, "malformed": 5, "malformation_details": Counter({"Pericardial oedema": 4, "Tail malformation": 2})},
            {"conc_id": "s3", "conc_type": "Substrate", "conc_value": 2.0, "total": 20, "dead": 9,
             "hatched": 8, "malformed": 10, "malformation_details": Counter({"Pericardial oedema": 8, "Tail malformation": 5})},
            {"conc_id": "s4", "conc_type": "Substrate", "conc_value": 4.0, "total": 20, "dead": 17,
             "hatched": 2, "malformed": 15, "malformation_details": Counter({"Pericardial oedema": 14, "Tail malformation": 11})},
            {"conc_id": "s5", "conc_type": "Substrate", "conc_value": 8.0, "total": 20, "dead": 20,
             "hatched": 0, "malformed": 18, "malformation_details": Counter({"Pericardial oedema": 17, "Tail malformation": 15})},
        ])
        stat = summary_df[summary_df.conc_type != "Positive Control"].copy()
        subs = stat[stat.conc_type == "Substrate"]
        plot = [{"type": "Substrate", "x": float(r.conc_value), "y": r.dead / r.total * 100} for _, r in subs.iterrows()]
        lc50 = calculate_lc50_robust(plot)
        lc50_num = float(lc50["_fitted_params"][3]) if lc50.get("_fitted_params") else None
        return {
            "lc50_results": lc50,
            "noec_loec_results": {"noec": "0.5000", "loec": "1.0000", "alpha_adjusted": 0.01},
            "sublethal_stats": sublethal_endpoint_stats(stat),
            "teratogenic_index": calculate_teratogenic_index(stat, lc50_num),
            "summary_df": summary_df,
            "mortality_plot_figure": None, "hatching_plot_figure": None,
            "malformation_plot_figure": None, "timecourse_plot_figure": None, "fate_plot_figure": None,
        }

    def test_sublethal_section_present(self, tmp_path, minimal_snapshot):
        out = str(tmp_path / "r.docx")
        ReportGenerator(minimal_snapshot, str(tmp_path), self._analysis()).generate_report(out)
        text = " ".join(p.text for p in Document(out).paragraphs)
        assert "Sublethal Endpoint Analysis" in text
        assert "Benjamini-Hochberg" in text

    def test_teratogenic_index_reported(self, tmp_path, minimal_snapshot):
        out = str(tmp_path / "r.docx")
        ReportGenerator(minimal_snapshot, str(tmp_path), self._analysis()).generate_report(out)
        text = " ".join(p.text for p in Document(out).paragraphs)
        assert "teratogenic index" in text.lower()

    def test_mortality_effect_size_table_present(self, tmp_path, minimal_snapshot):
        out = str(tmp_path / "r.docx")
        ReportGenerator(minimal_snapshot, str(tmp_path), self._analysis()).generate_report(out)
        headers = []
        for table in Document(out).tables:
            headers.append([c.text for c in table.rows[0].cells])
        assert any("Odds ratio (95% CI)" in h for h in headers)

    def test_sublethal_or_table_has_bh_column(self, tmp_path, minimal_snapshot):
        out = str(tmp_path / "r.docx")
        ReportGenerator(minimal_snapshot, str(tmp_path), self._analysis()).generate_report(out)
        headers = [[c.text for c in t.rows[0].cells] for t in Document(out).tables]
        assert any("p (BH)" in h for h in headers)


class TestReportSectionNumberingAndGating:
    """Regressions from the blind review: contiguous 2.x numbering and gated Methods claims."""

    def _analysis_sublethal_no_figs(self):
        from collections import Counter
        from src.core.biostatistics import sublethal_endpoint_stats
        summary_df = pd.DataFrame([
            {"conc_id": "ctrl", "conc_type": "Control",  "conc_value": 0.0, "total": 20, "dead": 1,
             "hatched": 16, "malformed": 0, "malformation_details": Counter()},
            {"conc_id": "s1", "conc_type": "Substrate", "conc_value": 1.0, "total": 20, "dead": 4,
             "hatched": 13, "malformed": 5, "malformation_details": Counter({"Pericardial oedema": 5})},
            {"conc_id": "s2", "conc_type": "Substrate", "conc_value": 2.0, "total": 20, "dead": 12,
             "hatched": 5, "malformed": 10, "malformation_details": Counter({"Pericardial oedema": 10})},
        ])
        return {
            "lc50_results": {"lc50": "Not Calculated", "model_info": {"display_name": "4PL", "mode": "manual", "n_free": 4}},
            "noec_loec_results": {"noec": "1.0000", "loec": "2.0000"},
            "trend_results": {"p_value": "Not Calculated", "trend": "Not Calculated"},
            "sublethal_stats": sublethal_endpoint_stats(summary_df),
            "summary_df": summary_df,
            "mortality_plot_figure": None, "hatching_plot_figure": None,
            "malformation_plot_figure": None, "timecourse_plot_figure": None, "fate_plot_figure": None,
        }

    def test_subsection_numbers_are_contiguous(self, tmp_path, minimal_snapshot):
        import re
        out = str(tmp_path / "r.docx")
        ReportGenerator(minimal_snapshot, str(tmp_path), self._analysis_sublethal_no_figs()).generate_report(out)
        headings = [p.text for p in Document(out).paragraphs
                    if "Heading" in p.style.name and re.match(r"^2\.\d", p.text.strip())]
        minor = sorted(int(re.match(r"^2\.(\d+)", h.strip()).group(1)) for h in headings)
        assert minor == list(range(1, len(minor) + 1)), f"non-contiguous: {headings}"
        # With no figures, the sublethal section immediately follows 2.2 -> 2.3
        assert any(h.startswith("2.3") and "Sublethal" in h for h in headings)

    def test_methods_omits_trend_and_effect_claims_without_control(self, tmp_path, minimal_snapshot, minimal_analysis):
        out = str(tmp_path / "r.docx")
        ReportGenerator(minimal_snapshot, str(tmp_path), minimal_analysis).generate_report(out)
        text = " ".join(p.text for p in Document(out).paragraphs)
        assert "Cochran-Armitage test for trend was additionally applied" not in text
        assert "Effect sizes are reported as odds ratios" not in text


class TestSurvivorMalformationDenominator:
    """Report 'Malformed (%)' uses survivors (live) as the denominator when present."""

    def test_report_malformed_pct_over_survivors(self, tmp_path, minimal_snapshot):
        summary_df = pd.DataFrame([
            {"conc_id": "ctrl", "conc_type": "Control",   "conc_value": 0.0, "total": 20, "dead": 0,
             "live": 20, "hatched": 18, "malformed": 0},
            {"conc_id": "s1",   "conc_type": "Substrate", "conc_value": 1.0, "total": 20, "dead": 12,
             "live": 8, "hatched": 6, "malformed": 6},  # 6/8 survivors = 75%, not 6/20 = 30%
        ])
        analysis = {
            "lc50_results": {"lc50": "Not Calculated", "model_info": {"display_name": "4PL", "mode": "manual", "n_free": 4}},
            "noec_loec_results": {"noec": "NC", "loec": "NC"},
            "summary_df": summary_df,
            "mortality_plot_figure": None, "hatching_plot_figure": None, "malformation_plot_figure": None,
        }
        out = str(tmp_path / "r.docx")
        ReportGenerator(minimal_snapshot, str(tmp_path), analysis).generate_report(out)
        # Find the summary table row for s1 and read its Malformed (%) cell
        malformed_vals = []
        for table in Document(out).tables:
            headers = [c.text for c in table.rows[0].cells]
            if "Malformed (%)" in headers:
                mi = headers.index("Malformed (%)")
                for row in table.rows[1:]:
                    malformed_vals.append(row.cells[mi].text)
        assert "75.00" in malformed_vals
        assert "30.00" not in malformed_vals


# ---------------------------------------------------------------------------
# The report must never assert more than the data supports
# ---------------------------------------------------------------------------

def _report_text(tmp_path, snapshot, analysis, name="claims.docx"):
    out = str(tmp_path / name)
    ReportGenerator(snapshot, str(tmp_path), analysis).generate_report(out)
    return " ".join(p.text for p in Document(out).paragraphs)


def _summary_frame(rows):
    return pd.DataFrame(rows)


class TestReportNeverRendersFailureMessagesAsResults:
    """Regression: an LC50 failure message was formatted as though it were a value.

    "100% mortality at all concentrations..." begins with a digit, so the old
    leading-character test classified it as numeric and produced
    "the LC50 was determined to be 100% mortality at all concentrations; ... mg/L
    (slope: Not Calculated)".
    """

    def _analysis(self, lc50_message):
        return {
            "lc50_results": {"lc50": lc50_message, "slope": "Not Calculated",
                             "r_squared": "Not Calculated", "model_info": {}},
            "noec_loec_results": {"noec": "16.0000", "noec_numeric": 16.0,
                                  "loec": "Not detected", "loec_numeric": None},
            "summary_df": _summary_frame([
                {"conc_id": "ctrl", "conc_type": "Control", "conc_value": 0.0,
                 "total": 20, "n_scored": 20, "dead": 20, "hatched": 0,
                 "live": 0, "malformed": 0},
                {"conc_id": "s1", "conc_type": "Substrate", "conc_value": 1.0,
                 "total": 20, "n_scored": 20, "dead": 20, "hatched": 0,
                 "live": 0, "malformed": 0},
            ]),
        }

    def test_hundred_percent_mortality_message_is_not_treated_as_a_value(
            self, tmp_path, minimal_snapshot):
        message = ("100% mortality at all concentrations; LC50 is below the "
                   "lowest concentration tested.")
        text = _report_text(tmp_path, minimal_snapshot, self._analysis(message))
        assert "the LC50 was determined to be 100% mortality" not in text
        assert "lowest concentration tested. mg/L" not in text
        assert f"LC50 calculation: {message}" in text
        # The message already ends in a full stop; the template must not add a
        # second one ("...concentration tested.. The NOEC was").
        assert "tested.. " not in text

    def test_non_numeric_noec_does_not_receive_a_unit(self, tmp_path, minimal_snapshot):
        analysis = self._analysis("Curve fitting failed; data may be inconsistent.")
        analysis["noec_loec_results"] = {"noec": "Control group not found",
                                         "noec_numeric": None,
                                         "loec": "Control group not found",
                                         "loec_numeric": None}
        text = _report_text(tmp_path, minimal_snapshot, analysis)
        assert "Control group not found mg/L" not in text
        assert "was Control group not found" in text


class TestSublethalSectionStatesOnlyWhatRan:
    def _analysis(self, sublethal_stats):
        return {
            "lc50_results": {"lc50": "Not Calculated", "model_info": {}},
            "noec_loec_results": {"noec": "NC", "loec": "NC"},
            "sublethal_stats": sublethal_stats,
            "summary_df": _summary_frame([
                {"conc_id": "ctrl", "conc_type": "Control", "conc_value": 0.0,
                 "total": 20, "n_scored": 20, "dead": 1, "hatched": 19,
                 "live": 19, "malformed": 0},
                {"conc_id": "s1", "conc_type": "Substrate", "conc_value": 1.0,
                 "total": 20, "n_scored": 20, "dead": 2, "hatched": 18,
                 "live": 18, "malformed": 0},
            ]),
        }

    def test_no_endpoints_scored_emits_no_sublethal_section(self, tmp_path, minimal_snapshot):
        """'available' is true via the pooled endpoint alone; the section is not."""
        stats = {"available": True, "tests": [], "ca": {},
                 "pooled": {"noec": "1.0000", "loec": "Not detected",
                            "noec_numeric": 1.0, "loec_numeric": None,
                            "no_events": True, "tests": [{"k": 0, "n": 18}]}}
        text = _report_text(tmp_path, minimal_snapshot, self._analysis(stats))
        assert "Sublethal Endpoint Analysis" not in text
        assert "Benjamini-Hochberg control of the false discovery rate" not in text

    def test_no_abnormalities_states_that_rather_than_a_noec(self, tmp_path, minimal_snapshot):
        stats = {"available": True, "tests": [], "ca": {},
                 "pooled": {"noec": "1.0000", "loec": "Not detected",
                            "noec_numeric": 1.0, "loec_numeric": None,
                            "no_events": True, "tests": [{"k": 0, "n": 18}]}}
        text = _report_text(tmp_path, minimal_snapshot, self._analysis(stats))
        assert "No morphological abnormalities were recorded in any group" in text
        assert "the sublethal NOEC and LOEC were" not in text


class TestObservationTimepointsFollowTheTest:
    """The report used to hard-code 96 hpf in twelve places."""

    def _analysis(self, analysis_day):
        return {
            "lc50_results": {"lc50": "Not Calculated", "model_info": {}},
            "noec_loec_results": {"noec": "NC", "loec": "NC"},
            "analysis_day": analysis_day,
            "summary_df": _summary_frame([
                {"conc_id": "ctrl", "conc_type": "Control", "conc_value": 0.0,
                 "total": 20, "n_scored": 20, "dead": 1, "hatched": 19,
                 "live": 19, "malformed": 0},
            ]),
        }

    def test_three_day_test_reports_seventy_two_hours(self, tmp_path, minimal_snapshot):
        snapshot = dict(minimal_snapshot, num_days=3)
        text = _report_text(tmp_path, snapshot, self._analysis(3))
        assert "72 hours post-fertilization" in text
        assert "96 hours post-fertilization" not in text

    def test_four_day_test_still_reports_ninety_six_hours(self, tmp_path, minimal_snapshot):
        text = _report_text(tmp_path, minimal_snapshot, self._analysis(4))
        assert "96 hours post-fertilization" in text

    def test_falls_back_to_project_duration_without_an_analysis_day(
            self, tmp_path, minimal_snapshot):
        """An older results dict must not break report generation."""
        analysis = self._analysis(None)
        analysis.pop("analysis_day")
        assert _report_text(tmp_path, minimal_snapshot, analysis)


class TestReferenceControlIsRecorded:
    """Which control the endpoints refer to is part of a reproducible report."""

    def _analysis(self, control_mode, comparison=None):
        return {
            "lc50_results": {"lc50": "Not Calculated", "model_info": {}},
            "noec_loec_results": {"noec": "NC", "loec": "NC"},
            "control_mode": control_mode,
            "control_comparison": comparison or {"applicable": False},
            "summary_df": _summary_frame([
                {"conc_id": "ctrl", "conc_type": "Control", "conc_value": 0.0,
                 "total": 20, "n_scored": 20, "dead": 1, "hatched": 19,
                 "live": 19, "malformed": 0},
                {"conc_id": "sc", "conc_type": "Solvent Control", "conc_value": 0.0,
                 "total": 20, "n_scored": 20, "dead": 9, "hatched": 11,
                 "live": 11, "malformed": 0},
            ]),
        }

    def test_pooled_reference_is_named(self, tmp_path, minimal_snapshot):
        text = _report_text(tmp_path, minimal_snapshot, self._analysis("pooled"))
        assert "against the pooled control" in text

    def test_solvent_reference_is_named(self, tmp_path, minimal_snapshot):
        text = _report_text(tmp_path, minimal_snapshot, self._analysis("solvent"))
        assert "against the solvent control" in text

    def test_control_comparison_is_reported_when_available(self, tmp_path, minimal_snapshot):
        comparison = {"applicable": True, "p_value": 0.0084, "differ": True,
                      "negative": {"dead": 1, "n": 20, "pct": 5.0},
                      "solvent": {"dead": 9, "n": 20, "pct": 45.0}}
        text = _report_text(tmp_path, minimal_snapshot,
                            self._analysis("control", comparison))
        assert "differed significantly" in text
        assert "two-sided" in text


class TestFertilizationRateCriterion:
    """OECD TG 236 §9a — the batch fertilization rate must be at least 70%.

    Optional: a project that did not record it is not a project that failed the
    criterion, so a blank field produces no statement either way.
    """

    def _msg(self, rate=None):
        conditions = {"acceptable_mortality": 10.0}
        if rate is not None:
            conditions["fertilization_rate"] = rate
        controls = pd.DataFrame([{
            "conc_id": "ctrl", "conc_type": "Control", "conc_value": 0.0,
            "total": 20, "n_scored": 20, "dead": 1, "hatched": 18, "live": 19,
            "malformed": 0,
        }])
        return ReportGenerator.get_test_validity_message(
            controls, {"test_conditions": conditions}, end_of_test=True, timepoint="96 hpf"
        )

    def test_an_unrecorded_rate_is_not_mentioned(self):
        assert "fertilization" not in self._msg().lower()

    def test_a_blank_rate_is_treated_as_unrecorded(self):
        assert "fertilization" not in self._msg("").lower()

    def test_a_rate_at_the_threshold_passes(self):
        msg = self._msg("70")
        assert "INVALID" not in msg
        assert "fertilization rate was 70%" in msg

    def test_a_rate_below_the_threshold_invalidates(self):
        msg = self._msg("55 %")
        assert "INVALID" in msg
        assert "below the minimum of 70%" in msg

    def test_the_unit_suffix_is_tolerated(self):
        """The field is free text, so '88 %' and '88' must read alike."""
        assert "88%" in self._msg("88 %")
        assert "88%" in self._msg("88")

    def test_it_is_assessed_before_the_end_of_the_test(self):
        """A batch property is known at plating, unlike hatching or the positive control."""
        controls = pd.DataFrame([{
            "conc_id": "ctrl", "conc_type": "Control", "conc_value": 0.0,
            "total": 20, "n_scored": 20, "dead": 1, "hatched": 0, "live": 19,
            "malformed": 0,
        }])
        msg = ReportGenerator.get_test_validity_message(
            controls, {"test_conditions": {"fertilization_rate": "88"}}, end_of_test=False
        )
        assert "fertilization rate was 88%" in msg
        assert "hatching" not in msg.lower()


class TestPositiveControlCriterion:
    """OECD TG 236 §9d — the positive control must reach 30% mortality at 96 h."""

    CTRL = {"conc_id": "ctrl", "conc_type": "Control", "conc_value": 0.0,
            "total": 20, "n_scored": 20, "dead": 1, "hatched": 18, "live": 19,
            "malformed": 0}

    def _msg(self, pc_dead=None, end_of_test=True):
        rows = [self.CTRL]
        if pc_dead is not None:
            rows.append({"conc_id": "pc", "conc_type": "Positive Control",
                         "conc_value": 4.0, "total": 20, "n_scored": 20,
                         "dead": pc_dead, "hatched": 5, "live": 20 - pc_dead,
                         "malformed": 0})
        return ReportGenerator.get_test_validity_message(
            pd.DataFrame(rows), {"test_conditions": {"acceptable_mortality": 10.0}},
            end_of_test=end_of_test, timepoint="96 hpf",
        )

    def test_a_study_without_a_positive_control_never_mentions_one(self):
        """Running without one is a design choice, not an unmet criterion."""
        assert "positive control" not in self._msg().lower()

    def test_a_responsive_positive_control_passes(self):
        msg = self._msg(pc_dead=12)
        assert "INVALID" not in msg
        assert "positive control at 96 hpf was 60.00%" in msg

    def test_at_the_threshold_it_passes(self):
        assert "INVALID" not in self._msg(pc_dead=6)

    def test_an_unresponsive_positive_control_invalidates(self):
        msg = self._msg(pc_dead=2)
        assert "INVALID" in msg
        assert "below the minimum of 30%" in msg

    def test_it_is_not_judged_before_the_end_of_the_test(self):
        assert "positive control" not in self._msg(pc_dead=2, end_of_test=False).lower()
